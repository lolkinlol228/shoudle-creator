from flask import Flask, request, jsonify, send_file, render_template_string
import sqlite3
import json
import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from contextlib import contextmanager
import threading

# ==================== CONFIGURATION ====================

class Config:
    DATABASE = 'schedule.db'
    EXPORT_FOLDER = 'exports'
    SECRET_KEY = 'your-secret-key-change-in-production'
    JSON_AS_ASCII = False

app = Flask(__name__)
app.config.from_object(Config)

for lang in ['russian', 'kyrgyz', 'english']:
    os.makedirs(os.path.join(app.config['EXPORT_FOLDER'], lang), exist_ok=True)

local = threading.local()

# ==================== DATABASE ====================

@contextmanager
def get_db():
    if not hasattr(local, 'db') or local.db is None:
        local.db = sqlite3.connect(
            app.config['DATABASE'],
            check_same_thread=False,
            timeout=10.0
        )
        local.db.row_factory = sqlite3.Row
        local.db.execute("PRAGMA foreign_keys = ON")
    
    try:
        yield local.db
    except Exception as e:
        local.db.rollback()
        raise e


def close_db_connection():
    if hasattr(local, 'db') and local.db is not None:
        local.db.close()
        local.db = None


@app.teardown_appcontext
def teardown_db(exception):
    close_db_connection()


def init_db():
    """Initialize database with all required tables"""
    with get_db() as conn:
        cursor = conn.cursor()
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS teachers (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name_russian TEXT NOT NULL,
                name_kyrgyz TEXT,
                name_english TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS subjects (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name_russian TEXT NOT NULL,
                name_kyrgyz TEXT,
                name_english TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS classrooms (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL UNIQUE,
                capacity INTEGER,
                building TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS groups (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL UNIQUE,
                year INTEGER,
                faculty TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS class_types (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name_russian TEXT NOT NULL,
                name_kyrgyz TEXT,
                name_english TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        cursor.execute("""
      CREATE TABLE IF NOT EXISTS system_settings (
    id INTEGER PRIMARY KEY CHECK (id = 1),
    class_duration_minutes INTEGER DEFAULT 80,
    default_break_minutes INTEGER DEFAULT 10,
    start_time_1 TEXT DEFAULT '08:00',
    start_time_2 TEXT DEFAULT '09:30',
    start_time_3 TEXT DEFAULT '11:10',
    start_time_4 TEXT DEFAULT '12:50',
    start_time_5 TEXT DEFAULT '14:30',
    start_time_6 TEXT DEFAULT '16:10',
    start_time_7 TEXT DEFAULT '17:50',
    start_time_8 TEXT DEFAULT '19:30',
    start_time_9 TEXT DEFAULT '21:10',
    start_time_10 TEXT DEFAULT '22:50',
    current_week_type TEXT DEFAULT 'numerator',
    week_start_date DATE NOT NULL,
    working_days_per_week INTEGER DEFAULT 6,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
)
        """)
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS schedules (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                date DATE NOT NULL,
                day_of_week INTEGER NOT NULL,
                start_time TIME NOT NULL,
                end_time TIME NOT NULL,
                teacher_id INTEGER NOT NULL,
                subject_id INTEGER NOT NULL,
                classroom_id INTEGER NOT NULL,
                class_type_id INTEGER NOT NULL,
                is_alternating BOOLEAN DEFAULT 0,
                week_type TEXT,
                break_after_minutes INTEGER,
                notes TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (teacher_id) REFERENCES teachers(id) ON DELETE RESTRICT,
                FOREIGN KEY (subject_id) REFERENCES subjects(id) ON DELETE RESTRICT,
                FOREIGN KEY (classroom_id) REFERENCES classrooms(id) ON DELETE RESTRICT,
                FOREIGN KEY (class_type_id) REFERENCES class_types(id) ON DELETE RESTRICT
            )
        """)
        
        cursor.execute("""CREATE INDEX IF NOT EXISTS idx_schedules_date ON schedules(date)""")
        cursor.execute("""CREATE INDEX IF NOT EXISTS idx_schedules_teacher ON schedules(teacher_id, date, start_time)""")
        cursor.execute("""CREATE INDEX IF NOT EXISTS idx_schedules_classroom ON schedules(classroom_id, date, start_time)""")
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS schedule_groups (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                schedule_id INTEGER NOT NULL,
                group_id INTEGER NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (schedule_id) REFERENCES schedules(id) ON DELETE CASCADE,
                FOREIGN KEY (group_id) REFERENCES groups(id) ON DELETE RESTRICT,
                UNIQUE(schedule_id, group_id)
            )
        """)
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS templates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT NOT NULL,
                description TEXT,
                source_date DATE,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS template_items (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                template_id INTEGER NOT NULL,
                day_of_week INTEGER NOT NULL,
                start_time TIME NOT NULL,
                end_time TIME NOT NULL,
                teacher_id INTEGER NOT NULL,
                subject_id INTEGER NOT NULL,
                classroom_id INTEGER NOT NULL,
                class_type_id INTEGER NOT NULL,
                is_alternating BOOLEAN DEFAULT 0,
                week_type TEXT,
                break_after_minutes INTEGER,
                groups_json TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (template_id) REFERENCES templates(id) ON DELETE CASCADE
            )
        """)
        
        cursor.execute("""
            INSERT OR IGNORE INTO system_settings 
            (id, week_start_date, current_week_type) 
            VALUES (1, date('now'), 'numerator')
        """)
        
 
        default_class_types = [
            ('Лекция', 'Лекция', 'Lecture'),
            ('Практическое занятие', 'Практикалык сабак', 'Practical'),
            ('Лабораторная работа', 'Лабораториялык иш', 'Laboratory'),
            ('Семинар', 'Семинар', 'Seminar'),
            ('Консультация', 'Консультация', 'Consultation')
        ]
        
        for ct in default_class_types:
            cursor.execute("""
                INSERT INTO class_types (name_russian, name_kyrgyz, name_english)
                SELECT ?, ?, ?
                WHERE NOT EXISTS (
                    SELECT 1 FROM class_types WHERE name_russian = ?
                )
            """, (ct[0], ct[1], ct[2], ct[0]))
        
        conn.commit()


init_db()




def calculate_week_type(target_date):
    with get_db() as conn:
        cursor = conn.cursor()
        settings = cursor.execute(
            "SELECT week_start_date, current_week_type FROM system_settings WHERE id = 1"
        ).fetchone()
        
        if not settings:
            return 'numerator'
        
        start_date = datetime.strptime(settings['week_start_date'], '%Y-%m-%d').date()
        if isinstance(target_date, str):
            target = datetime.strptime(target_date, '%Y-%m-%d').date()
        else:
            target = target_date
        
        start_week_type = settings['current_week_type']
        days_diff = (target - start_date).days
        weeks_diff = days_diff // 7
        
        if weeks_diff % 2 == 0:
            return start_week_type
        else:
            return 'denominator' if start_week_type == 'numerator' else 'numerator'

def calculate_end_time(start_time_str, duration_minutes):
    start = datetime.strptime(start_time_str, '%H:%M')
    end = start + timedelta(minutes=duration_minutes)
    return end.strftime('%H:%M')

def check_conflicts(schedule_data, exclude_schedule_id=None):
    conflicts = []
    
    with get_db() as conn:
        cursor = conn.cursor()
        
        date = schedule_data['date']
        start_time = schedule_data['start_time']
        end_time = schedule_data['end_time']
        teacher_id = schedule_data['teacher_id']
        classroom_id = schedule_data['classroom_id']
        week_type = schedule_data.get('week_type')
        is_alternating = schedule_data.get('is_alternating', 0)
        
        query = """
            SELECT s.*, 
                   t.name_russian as teacher_name,
                   sub.name_russian as subject_name,
                   c.name as classroom_name,
                   GROUP_CONCAT(g.name) as group_names
            FROM schedules s
            JOIN teachers t ON s.teacher_id = t.id
            JOIN subjects sub ON s.subject_id = sub.id
            JOIN classrooms c ON s.classroom_id = c.id
            LEFT JOIN schedule_groups sg ON s.id = sg.schedule_id
            LEFT JOIN groups g ON sg.group_id = g.id
            WHERE s.date = ?
              AND s.id != ?
              AND (
                  (s.start_time < ? AND s.end_time > ?)
                  OR (s.start_time >= ? AND s.start_time < ?)
                  OR (s.end_time > ? AND s.end_time <= ?)
              )
            GROUP BY s.id
        """
        
        exclude_id = exclude_schedule_id or -1
        params = (date, exclude_id, end_time, start_time, 
                 start_time, end_time, start_time, end_time)
        
        overlapping = cursor.execute(query, params).fetchall()
        
        for existing in overlapping:
            if is_alternating and existing['is_alternating']:
                if week_type and existing['week_type']:
                    if week_type != existing['week_type']:
                        continue
            
            if existing['teacher_id'] == teacher_id:
                conflicts.append({
                    'type': 'teacher',
                    'message': f"Преподаватель {existing['teacher_name']} занят в группе {existing['group_names']} "
                               f"с {existing['start_time']} до {existing['end_time']} "
                               f"(предмет: {existing['subject_name']})",
                    'schedule_id': existing['id']
                })
            
            if existing['classroom_id'] == classroom_id:
                conflicts.append({
                    'type': 'classroom',
                    'message': f"Аудитория {existing['classroom_name']} занята группой {existing['group_names']} "
                               f"с {existing['start_time']} до {existing['end_time']} "
                               f"(преподаватель: {existing['teacher_name']}, предмет: {existing['subject_name']})",
                    'schedule_id': existing['id']
                })
    
    return conflicts



@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/api/week_type', methods=['GET'])
def get_week_type():
    date = request.args.get('date', datetime.now().strftime('%Y-%m-%d'))
    week_type = calculate_week_type(date)
    return jsonify({'date': date, 'week_type': week_type})

@app.route('/api/settings', methods=['GET'])
def get_settings():
    with get_db() as conn:
        cursor = conn.cursor()
        settings = cursor.execute("SELECT * FROM system_settings WHERE id = 1").fetchone()
        return jsonify(dict(settings) if settings else {})

@app.route('/api/settings', methods=['PUT'])
def update_settings():
    data = request.json
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE system_settings SET
                class_duration_minutes = ?,
                default_break_minutes = ?,
                start_time_1 = ?,
                start_time_2 = ?,
                start_time_3 = ?,
                start_time_4 = ?,
                start_time_5 = ?,
                start_time_6 = ?,
                start_time_7 = ?,
                start_time_8 = ?,
                start_time_9 = ?,
                start_time_10 = ?,
                current_week_type = ?,
                week_start_date = ?,
                working_days_per_week = ?,
                updated_at = CURRENT_TIMESTAMP
            WHERE id = 1
        """, (
            data.get('class_duration_minutes'),
            data.get('default_break_minutes'),
            data.get('start_time_1'),
            data.get('start_time_2'),
            data.get('start_time_3'),
            data.get('start_time_4'),
            data.get('start_time_5'),
            data.get('start_time_6'),
            data.get('start_time_7'),
            data.get('start_time_8'),
            data.get('start_time_9'),
            data.get('start_time_10'),
            data.get('current_week_type'),
            data.get('week_start_date'),
            data.get('working_days_per_week')
        ))
        conn.commit()
        return jsonify({'success': True})

@app.route('/api/teachers', methods=['GET'])
def get_teachers():
    with get_db() as conn:
        cursor = conn.cursor()
        teachers = cursor.execute("SELECT * FROM teachers ORDER BY name_russian").fetchall()
        return jsonify([dict(t) for t in teachers])

@app.route('/api/teachers', methods=['POST'])
def create_teacher():
    data = request.json
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO teachers (name_russian, name_kyrgyz, name_english)
            VALUES (?, ?, ?)
        """, (data['name_russian'], data.get('name_kyrgyz'), data.get('name_english')))
        conn.commit()
        return jsonify({'id': cursor.lastrowid, 'success': True})

@app.route('/api/teachers/<int:id>', methods=['GET'])
def get_teacher(id):
    with get_db() as conn:
        cursor = conn.cursor()
        teacher = cursor.execute("SELECT * FROM teachers WHERE id = ?", (id,)).fetchone()
        return jsonify(dict(teacher) if teacher else {})

@app.route('/api/teachers/<int:id>', methods=['PUT'])
def update_teacher(id):
    data = request.json
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE teachers SET name_russian = ?, name_kyrgyz = ?, name_english = ?
            WHERE id = ?
        """, (data['name_russian'], data.get('name_kyrgyz'), data.get('name_english'), id))
        conn.commit()
        return jsonify({'success': True})

@app.route('/api/teachers/<int:id>', methods=['DELETE'])
def delete_teacher(id):
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM teachers WHERE id = ?", (id,))
            conn.commit()
            return jsonify({'success': True})
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': 'Невозможно удалить преподавателя, используется в расписании'}), 400

@app.route('/api/subjects', methods=['GET'])
def get_subjects():
    with get_db() as conn:
        cursor = conn.cursor()
        subjects = cursor.execute("SELECT * FROM subjects ORDER BY name_russian").fetchall()
        return jsonify([dict(s) for s in subjects])

@app.route('/api/subjects', methods=['POST'])
def create_subject():
    data = request.json
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO subjects (name_russian, name_kyrgyz, name_english)
            VALUES (?, ?, ?)
        """, (data['name_russian'], data.get('name_kyrgyz'), data.get('name_english')))
        conn.commit()
        return jsonify({'id': cursor.lastrowid, 'success': True})

@app.route('/api/subjects/<int:id>', methods=['GET'])
def get_subject(id):
    with get_db() as conn:
        cursor = conn.cursor()
        subject = cursor.execute("SELECT * FROM subjects WHERE id = ?", (id,)).fetchone()
        return jsonify(dict(subject) if subject else {})

@app.route('/api/subjects/<int:id>', methods=['PUT'])
def update_subject(id):
    data = request.json
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE subjects SET name_russian = ?, name_kyrgyz = ?, name_english = ?
            WHERE id = ?
        """, (data['name_russian'], data.get('name_kyrgyz'), data.get('name_english'), id))
        conn.commit()
        return jsonify({'success': True})

@app.route('/api/subjects/<int:id>', methods=['DELETE'])
def delete_subject(id):
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM subjects WHERE id = ?", (id,))
            conn.commit()
            return jsonify({'success': True})
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': 'Невозможно удалить предмет, используется в расписании'}), 400


@app.route('/api/classrooms', methods=['GET'])
def get_classrooms():
    with get_db() as conn:
        cursor = conn.cursor()
        classrooms = cursor.execute("SELECT * FROM classrooms ORDER BY name").fetchall()
        return jsonify([dict(c) for c in classrooms])

@app.route('/api/classrooms', methods=['POST'])
def create_classroom():
    data = request.json
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO classrooms (name, capacity, building)
            VALUES (?, ?, ?)
        """, (data['name'], data.get('capacity'), data.get('building')))
        conn.commit()
        return jsonify({'id': cursor.lastrowid, 'success': True})

@app.route('/api/classrooms/<int:id>', methods=['GET'])
def get_classroom(id):
    with get_db() as conn:
        cursor = conn.cursor()
        classroom = cursor.execute("SELECT * FROM classrooms WHERE id = ?", (id,)).fetchone()
        return jsonify(dict(classroom) if classroom else {})

@app.route('/api/classrooms/<int:id>', methods=['PUT'])
def update_classroom(id):
    data = request.json
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE classrooms SET name = ?, capacity = ?, building = ?
            WHERE id = ?
        """, (data['name'], data.get('capacity'), data.get('building'), id))
        conn.commit()
        return jsonify({'success': True})

@app.route('/api/classrooms/<int:id>', methods=['DELETE'])
def delete_classroom(id):
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM classrooms WHERE id = ?", (id,))
            conn.commit()
            return jsonify({'success': True})
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': 'Невозможно удалить аудиторию, используется в расписании'}), 400


@app.route('/api/groups', methods=['GET'])
def get_groups():
    with get_db() as conn:
        cursor = conn.cursor()
        groups = cursor.execute("SELECT * FROM groups ORDER BY name").fetchall()
        return jsonify([dict(g) for g in groups])

@app.route('/api/groups', methods=['POST'])
def create_group():
    data = request.json
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO groups (name, year, faculty)
            VALUES (?, ?, ?)
        """, (data['name'], data.get('year'), data.get('faculty')))
        conn.commit()
        return jsonify({'id': cursor.lastrowid, 'success': True})

@app.route('/api/groups/<int:id>', methods=['GET'])
def get_group(id):
    with get_db() as conn:
        cursor = conn.cursor()
        group = cursor.execute("SELECT * FROM groups WHERE id = ?", (id,)).fetchone()
        return jsonify(dict(group) if group else {})

@app.route('/api/groups/<int:id>', methods=['PUT'])
def update_group(id):
    data = request.json
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE groups SET name = ?, year = ?, faculty = ?
            WHERE id = ?
        """, (data['name'], data.get('year'), data.get('faculty'), id))
        conn.commit()
        return jsonify({'success': True})

@app.route('/api/groups/<int:id>', methods=['DELETE'])
def delete_group(id):
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM groups WHERE id = ?", (id,))
            conn.commit()
            return jsonify({'success': True})
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': 'Невозможно удалить группу, используется в расписании'}), 400

@app.route('/api/class_types', methods=['GET'])
def get_class_types():
    with get_db() as conn:
        cursor = conn.cursor()
        class_types = cursor.execute("SELECT * FROM class_types ORDER BY name_russian").fetchall()
        return jsonify([dict(ct) for ct in class_types])

@app.route('/api/class_types', methods=['POST'])
def create_class_type():
    data = request.json
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO class_types (name_russian, name_kyrgyz, name_english)
            VALUES (?, ?, ?)
        """, (data['name_russian'], data.get('name_kyrgyz'), data.get('name_english')))
        conn.commit()
        return jsonify({'id': cursor.lastrowid, 'success': True})

@app.route('/api/class_types/<int:id>', methods=['GET'])
def get_class_type(id):
    with get_db() as conn:
        cursor = conn.cursor()
        class_type = cursor.execute("SELECT * FROM class_types WHERE id = ?", (id,)).fetchone()
        return jsonify(dict(class_type) if class_type else {})

@app.route('/api/class_types/<int:id>', methods=['PUT'])
def update_class_type(id):
    data = request.json
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE class_types SET name_russian = ?, name_kyrgyz = ?, name_english = ?
            WHERE id = ?
        """, (data['name_russian'], data.get('name_kyrgyz'), data.get('name_english'), id))
        conn.commit()
        return jsonify({'success': True})

@app.route('/api/class_types/<int:id>', methods=['DELETE'])
def delete_class_type(id):
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM class_types WHERE id = ?", (id,))
            conn.commit()
            return jsonify({'success': True})
    except sqlite3.IntegrityError:
        return jsonify({'success': False, 'error': 'Невозможно удалить тип занятия, используется в расписании'}), 400


@app.route('/api/schedules', methods=['GET'])
def get_schedules():
    date = request.args.get('date')
    group_id = request.args.get('group_id')
    teacher_id = request.args.get('teacher_id')
    date_from = request.args.get('date_from')
    date_to = request.args.get('date_to')
    
    with get_db() as conn:
        cursor = conn.cursor()
        
        query = """
            SELECT s.*, 
                   t.name_russian as teacher_name_ru,
                   t.name_kyrgyz as teacher_name_ky,
                   t.name_english as teacher_name_en,
                   sub.name_russian as subject_name_ru,
                   sub.name_kyrgyz as subject_name_ky,
                   sub.name_english as subject_name_en,
                   c.name as classroom_name,
                   ct.name_russian as class_type_ru,
                   ct.name_kyrgyz as class_type_ky,
                   ct.name_english as class_type_en,
                   GROUP_CONCAT(DISTINCT g.id) as group_ids,
                   GROUP_CONCAT(DISTINCT g.name) as group_names
            FROM schedules s
            JOIN teachers t ON s.teacher_id = t.id
            JOIN subjects sub ON s.subject_id = sub.id
            JOIN classrooms c ON s.classroom_id = c.id
            JOIN class_types ct ON s.class_type_id = ct.id
            LEFT JOIN schedule_groups sg ON s.id = sg.schedule_id
            LEFT JOIN groups g ON sg.group_id = g.id
            WHERE 1=1
        """
        
        params = []
        if date:
            query += " AND s.date = ?"
            params.append(date)
        if date_from:
            query += " AND s.date >= ?"
            params.append(date_from)
        if date_to:
            query += " AND s.date <= ?"
            params.append(date_to)
        if group_id:
            query += " AND sg.group_id = ?"
            params.append(group_id)
        if teacher_id:
            query += " AND s.teacher_id = ?"
            params.append(teacher_id)
        
        query += " GROUP BY s.id ORDER BY s.date, s.start_time"
        schedules = cursor.execute(query, params).fetchall()
        return jsonify([dict(schedule) for schedule in schedules])

@app.route('/api/schedules', methods=['POST'])
def create_schedule():
    data = request.json
    
    if not data.get('end_time'):
        with get_db() as conn:
            cursor = conn.cursor()
            settings = cursor.execute("SELECT class_duration_minutes FROM system_settings WHERE id = 1").fetchone()
            if settings:
                data['end_time'] = calculate_end_time(data['start_time'], settings['class_duration_minutes'])
    
    date_obj = datetime.strptime(data['date'], '%Y-%m-%d')
    data['day_of_week'] = date_obj.isoweekday()
    
    if data.get('is_alternating') and not data.get('week_type'):
        data['week_type'] = calculate_week_type(data['date'])
    
    conflicts = check_conflicts(data)
    if conflicts:
        return jsonify({'success': False, 'conflicts': conflicts}), 409
    
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO schedules 
            (date, day_of_week, start_time, end_time, teacher_id, subject_id, 
             classroom_id, class_type_id, is_alternating, week_type, break_after_minutes, notes)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            data['date'], data['day_of_week'], data['start_time'], data['end_time'],
            data['teacher_id'], data['subject_id'], data['classroom_id'], 
            data['class_type_id'], data.get('is_alternating', 0), 
            data.get('week_type'), data.get('break_after_minutes'), data.get('notes')
        ))
        
        schedule_id = cursor.lastrowid
        
        for group_id in data.get('group_ids', []):
            cursor.execute(
                "INSERT INTO schedule_groups (schedule_id, group_id) VALUES (?, ?)",
                (schedule_id, group_id)
            )
        
        conn.commit()
        return jsonify({'id': schedule_id, 'success': True})

@app.route('/api/schedules/<int:id>', methods=['GET'])
def get_schedule(id):
    with get_db() as conn:
        cursor = conn.cursor()
        schedule = cursor.execute("""
            SELECT s.*, GROUP_CONCAT(sg.group_id) as group_ids
            FROM schedules s
            LEFT JOIN schedule_groups sg ON s.id = sg.schedule_id
            WHERE s.id = ?
            GROUP BY s.id
        """, (id,)).fetchone()
        
        if schedule:
            result = dict(schedule)
            if result['group_ids']:
                result['group_ids'] = [int(gid) for gid in result['group_ids'].split(',')]
            else:
                result['group_ids'] = []
            return jsonify(result)
        return jsonify({})

@app.route('/api/schedules/<int:id>', methods=['PUT'])
def update_schedule(id):
    data = request.json
    data['id'] = id
    
    conflicts = check_conflicts(data, id)
    if conflicts:
        return jsonify({'success': False, 'conflicts': conflicts}), 409
    
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE schedules SET
                date = ?, start_time = ?, end_time = ?,
                teacher_id = ?, subject_id = ?, classroom_id = ?,
                class_type_id = ?, is_alternating = ?, week_type = ?,
                break_after_minutes = ?, notes = ?,
                updated_at = CURRENT_TIMESTAMP
            WHERE id = ?
        """, (
            data['date'], data['start_time'], data['end_time'],
            data['teacher_id'], data['subject_id'], data['classroom_id'],
            data['class_type_id'], data.get('is_alternating', 0),
            data.get('week_type'), data.get('break_after_minutes'),
            data.get('notes'), id
        ))
        
        cursor.execute("DELETE FROM schedule_groups WHERE schedule_id = ?", (id,))
        for group_id in data.get('group_ids', []):
            cursor.execute(
                "INSERT INTO schedule_groups (schedule_id, group_id) VALUES (?, ?)",
                (id, group_id)
            )
        
        conn.commit()
        return jsonify({'success': True})

@app.route('/api/schedules/<int:id>', methods=['DELETE'])
def delete_schedule(id):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM schedules WHERE id = ?", (id,))
        conn.commit()
        return jsonify({'success': True})

@app.route('/api/schedules/validate', methods=['POST'])
def validate_schedule():
    data = request.json
    conflicts = check_conflicts(data, data.get('id'))
    return jsonify({'valid': len(conflicts) == 0, 'conflicts': conflicts})


@app.route('/api/templates', methods=['GET'])
def get_templates():
    with get_db() as conn:
        cursor = conn.cursor()
        templates = cursor.execute("SELECT * FROM templates ORDER BY created_at DESC").fetchall()
        return jsonify([dict(t) for t in templates])

@app.route('/api/templates', methods=['POST'])
def create_template():
    data = request.json
    source_date = data.get('source_date')
    
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO templates (name, description, source_date) VALUES (?, ?, ?)",
            (data['name'], data.get('description'), source_date)
        )
        template_id = cursor.lastrowid
        
        schedules = cursor.execute("""
            SELECT s.*, GROUP_CONCAT(sg.group_id) as group_ids
            FROM schedules s
            LEFT JOIN schedule_groups sg ON s.id = sg.schedule_id
            WHERE s.date = ?
            GROUP BY s.id
        """, (source_date,)).fetchall()
        
        for schedule in schedules:
            cursor.execute("""
                INSERT INTO template_items
                (template_id, day_of_week, start_time, end_time, teacher_id, 
                 subject_id, classroom_id, class_type_id, is_alternating, 
                 week_type, break_after_minutes, groups_json)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                template_id, schedule['day_of_week'], schedule['start_time'],
                schedule['end_time'], schedule['teacher_id'], schedule['subject_id'],
                schedule['classroom_id'], schedule['class_type_id'], 
                schedule['is_alternating'], schedule['week_type'],
                schedule['break_after_minutes'], schedule['group_ids'] or ''
            ))
        
        conn.commit()
        return jsonify({'id': template_id, 'success': True})

@app.route('/api/templates/<int:id>/apply', methods=['POST'])
def apply_template(id):
    data = request.json
    target_dates = data.get('dates', [])
    date_range = data.get('date_range')
    
    if date_range:
        start = datetime.strptime(date_range['start'], '%Y-%m-%d')
        end = datetime.strptime(date_range['end'], '%Y-%m-%d')
        target_dates = []
        current = start
        while current <= end:
            target_dates.append(current.strftime('%Y-%m-%d'))
            current += timedelta(days=1)
    
    with get_db() as conn:
        cursor = conn.cursor()
        items = cursor.execute("SELECT * FROM template_items WHERE template_id = ?", (id,)).fetchall()
        
        created_count = 0
        conflicts_found = []
        
        for target_date in target_dates:
            date_obj = datetime.strptime(target_date, '%Y-%m-%d')
            day_of_week = date_obj.isoweekday()
            week_type = calculate_week_type(target_date)
            
            for item in items:
                group_ids_str = item['groups_json']
                group_ids = [int(gid) for gid in group_ids_str.split(',') if gid] if group_ids_str else []
                
                schedule_data = {
                    'date': target_date,
                    'day_of_week': day_of_week,
                    'start_time': item['start_time'],
                    'end_time': item['end_time'],
                    'teacher_id': item['teacher_id'],
                    'subject_id': item['subject_id'],
                    'classroom_id': item['classroom_id'],
                    'class_type_id': item['class_type_id'],
                    'is_alternating': item['is_alternating'],
                    'week_type': week_type if item['is_alternating'] else None,
                    'break_after_minutes': item['break_after_minutes'],
                    'group_ids': group_ids
                }
                
                conflicts = check_conflicts(schedule_data)
                if conflicts:
                    conflicts_found.append({
                        'date': target_date,
                        'time': f"{item['start_time']}-{item['end_time']}",
                        'conflicts': conflicts
                    })
                    continue
                
                cursor.execute("""
                    INSERT INTO schedules 
                    (date, day_of_week, start_time, end_time, teacher_id, subject_id, 
                     classroom_id, class_type_id, is_alternating, week_type, break_after_minutes)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    schedule_data['date'], schedule_data['day_of_week'],
                    schedule_data['start_time'], schedule_data['end_time'],
                    schedule_data['teacher_id'], schedule_data['subject_id'],
                    schedule_data['classroom_id'], schedule_data['class_type_id'],
                    schedule_data['is_alternating'], schedule_data['week_type'],
                    schedule_data['break_after_minutes']
                ))
                
                schedule_id = cursor.lastrowid
                
                for group_id in group_ids:
                    cursor.execute(
                        "INSERT INTO schedule_groups (schedule_id, group_id) VALUES (?, ?)",
                        (schedule_id, group_id)
                    )
                
                created_count += 1
        
        conn.commit()
        return jsonify({
            'success': True,
            'created_count': created_count,
            'conflicts': conflicts_found
        })

@app.route('/api/templates/<int:id>', methods=['DELETE'])
def delete_template(id):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM templates WHERE id = ?", (id,))
        conn.commit()
        return jsonify({'success': True})


def format_docx_table(table):
    table.style = 'Light Grid Accent 1'
    for cell in table.rows[0].cells:
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading_elm = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w')))
            cell._element.get_or_add_tcPr().append(shading_elm)
    
    for row in table.rows[1:]:
        for cell in row.cells:
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)


from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

@app.route('/api/export_docx', methods=['POST'])
def export_docx():
    data = request.json
    language = data.get('language', 'russian')
    report_type = data.get('report_type', 'general')
    filters = data.get('filters', {})
    
    # Извлечение фильтров
    group_id = filters.get('group_id')
    year = filters.get('year')
    teacher_id = filters.get('teacher_id')
    date_from = filters.get('date_from')
    date_to = filters.get('date_to')
    
    lang_suffix = {'russian': 'russian', 'kyrgyz': 'kyrgyz', 'english': 'english'}[language]

    with get_db() as conn:
        cursor = conn.cursor()
        
        # Базовый запрос
        query = f"""
            SELECT s.date, g.name as group_name, g.year,
                   CASE WHEN '{lang_suffix}' = 'russian' THEN t.name_russian
                        WHEN '{lang_suffix}' = 'kyrgyz' THEN t.name_kyrgyz
                        ELSE t.name_english END as teacher,
                   CASE WHEN '{lang_suffix}' = 'russian' THEN sub.name_russian
                        WHEN '{lang_suffix}' = 'kyrgyz' THEN sub.name_kyrgyz
                        ELSE sub.name_english END as subject,
                   CASE WHEN '{lang_suffix}' = 'russian' THEN ct.name_russian
                        WHEN '{lang_suffix}' = 'kyrgyz' THEN ct.name_kyrgyz
                        ELSE ct.name_english END as class_type,
                   s.start_time || ' - ' || s.end_time as time,
                   s.start_time,
                   c.name as classroom,
                   g.id as group_id,
                   s.day_of_week,
                   t.id as teacher_id
              FROM schedules s
              JOIN schedule_groups sg ON s.id=sg.schedule_id
              JOIN groups g ON sg.group_id=g.id
              JOIN teachers t ON s.teacher_id=t.id
              JOIN subjects sub ON s.subject_id=sub.id
              JOIN classrooms c ON s.classroom_id=c.id
              JOIN class_types ct ON s.class_type_id=ct.id
              WHERE 1=1
        """
        params = []
        
        # Фильтр по датам
        if date_from:
            query += " AND s.date >= ?"
            params.append(date_from)
        if date_to:
            query += " AND s.date <= ?"
            params.append(date_to)
        
        # Фильтры в зависимости от типа отчета
        if report_type == 'group':
            if group_id:
                query += " AND g.id = ?"
                params.append(group_id)
            elif year:
                query += " AND g.year = ?"
                params.append(year)
        elif report_type == 'teacher':
            if teacher_id:
                query += " AND t.id = ?"
                params.append(teacher_id)
        
        query += " ORDER BY g.name, s.day_of_week, s.start_time"
        rows = cursor.execute(query, params).fetchall()

    if not rows:
        return jsonify({'success': False, 'error': 'Нет данных для экспорта'}), 404

    from collections import defaultdict
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    def set_cell_background(cell, fill_color):
        """Установка цвета фона ячейки"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), fill_color)
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    def set_cell_vertical_alignment(cell, align="center"):
        """Установка вертикального выравнивания ячейки"""
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), align)
        tcPr.append(vAlign)
    
    # Названия дней недели
    day_names = {
        'russian': ['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота'],
        'kyrgyz': ['Дүйшөмбү', 'Шейшемби', 'Шаршемби', 'Бейшемби', 'Жума', 'Ишемби'],
        'english': ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    }[language]
    
    # Заголовки для разных типов отчетов
    headers = {
        'russian': {
            'general': 'Расписание занятий',
            'group': 'Расписание группы',
            'teacher': 'Расписание преподавателя',
            'period': 'Период',
            'group_label': 'Группа',
            'course': 'курс',
            'subject': 'Предмет',
            'type': 'Тип',
            'teacher_label': 'Преподаватель',
            'time': 'Время',
            'room': 'Аудитория'
        },
        'kyrgyz': {
            'general': 'Сабак расписаниеси',
            'group': 'Топтун расписаниеси',
            'teacher': 'Окутуучунун расписаниеси',
            'period': 'Мезгил',
            'group_label': 'Топ',
            'course': 'курс',
            'subject': 'Предмет',
            'type': 'Түрү',
            'teacher_label': 'Окутуучу',
            'time': 'Убакыт',
            'room': 'Аудитория'
        },
        'english': {
            'general': 'Class Schedule',
            'group': 'Group Schedule',
            'teacher': 'Teacher Schedule',
            'period': 'Period',
            'group_label': 'Group',
            'course': 'course',
            'subject': 'Subject',
            'type': 'Type',
            'teacher_label': 'Teacher',
            'time': 'Time',
            'room': 'Room'
        }
    }[language]
    
    doc = Document()
    
    # Заголовок документа
    title = doc.add_paragraph()
    if report_type == 'group':
        if group_id:
            title_text = headers['group']
        else:
            title_text = f"{headers['general']} ({year} {headers['course']})"
    elif report_type == 'teacher':
        title_text = headers['teacher']
    else:
        title_text = headers['general']
    
    title_run = title.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(68, 114, 196)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(10)
    
    # Период
    if date_from and date_to:
        period = doc.add_paragraph()
        period_run = period.add_run(f"{headers['period']}: {date_from} — {date_to}")
        period_run.font.size = Pt(11)
        period.alignment = WD_ALIGN_PARAGRAPH.CENTER
        period.space_after = Pt(20)
    
    # Группировка данных
    if report_type == 'teacher':
        # Для преподавателя группируем по дням
        data_by_day = defaultdict(list)
        teacher_name = rows[0]['teacher'] if rows else ''
        
        for r in rows:
            data_by_day[r['day_of_week']].append(r)
        
        # Добавляем имя преподавателя
        teacher_p = doc.add_paragraph()
        teacher_run = teacher_p.add_run(f"{headers['teacher_label']}: {teacher_name}")
        teacher_run.bold = True
        teacher_run.font.size = Pt(14)
        teacher_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        teacher_p.space_after = Pt(15)
        
        # Создаем таблицу по дням
        for day_num in range(1, 7):
            if day_num not in data_by_day:
                continue
                
            lessons = data_by_day[day_num]
            
            # День недели
            day_p = doc.add_paragraph()
            day_run = day_p.add_run(day_names[day_num - 1])
            day_run.bold = True
            day_run.font.size = Pt(12)
            day_p.space_before = Pt(10)
            day_p.space_after = Pt(5)
            
            # Таблица с занятиями
            table = doc.add_table(rows=len(lessons) + 1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            # Заголовки
            headers_row = table.rows[0].cells
            headers_row[0].text = headers['time']
            headers_row[1].text = headers['subject']
            headers_row[2].text = headers['type']
            headers_row[3].text = headers['group_label']
            headers_row[4].text = headers['room']
            
            for cell in headers_row:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Данные
            for i, lesson in enumerate(lessons, 1):
                row = table.rows[i].cells
                row[0].text = lesson['time']
                row[1].text = lesson['subject']
                row[2].text = lesson['class_type']
                row[3].text = lesson['group_name']
                row[4].text = lesson['classroom']
                
                for cell in row:
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    else:
        # Для групп - стандартная группировка
        data_by_group = defaultdict(lambda: defaultdict(list))
        for r in rows:
            data_by_group[r['group_name']][r['day_of_week']].append(r)
        
        for group_name, days in data_by_group.items():
            # Заголовок группы
            p = doc.add_paragraph()
            group_run = p.add_run(f"{headers['group_label']}: {group_name}")
            group_run.bold = True
            group_run.font.size = Pt(14)
            group_run.font.color.rgb = RGBColor(68, 114, 196)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_after = Pt(12)
            
            # Определение максимального количества пар в день
            max_lessons = max(len(lessons) for lessons in days.values()) if days else 1
            
            # Создание таблицы
            rows_count = 6
            table = doc.add_table(rows=rows_count, cols=max_lessons + 1)
            table.style = 'Table Grid'
            
            # Заполнение дней недели
            for i in range(rows_count):
                cell_day = table.cell(i, 0)
                cell_day.text = day_names[i]
                cell_day.width = Inches(1.2)
                
                # Центрирование текста по горизонтали и вертикали
                set_cell_vertical_alignment(cell_day, "center")
                set_cell_background(cell_day, "D9E2F3")  # Светло-голубой фон
                
                for paragraph in cell_day.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0, 0, 0)
            
            # Заполнение расписания
            for i in range(rows_count):
                lessons = days.get(i + 1, [])
                for j, lesson in enumerate(lessons):
                    cell_lesson = table.cell(i, j + 1)
                    cell_lesson.text = ""  # Очищаем ячейку
                    
                    # Центрирование по вертикали
                    set_cell_vertical_alignment(cell_lesson, "center")
                    
                    # Добавляем информацию построчно с форматированием
                    p1 = cell_lesson.paragraphs[0]
                    
                    # Предмет (жирным шрифтом, центрирован)
                    run_subject = p1.add_run(f"{lesson['subject']}\n")
                    run_subject.font.bold = True
                    run_subject.font.size = Pt(10)
                    run_subject.font.color.rgb = RGBColor(0, 0, 128)
                    
                    # Тип занятия
                    run_type = p1.add_run(f"{headers['type']}: {lesson['class_type']}\n")
                    run_type.font.size = Pt(9)
                    
                    # Разделитель
                    run_div1 = p1.add_run("─────────────\n")
                    run_div1.font.size = Pt(8)
                    run_div1.font.color.rgb = RGBColor(150, 150, 150)
                    
                    # Преподаватель
                    run_teacher_label = p1.add_run(f"{headers['teacher_label']}: ")
                    run_teacher_label.font.size = Pt(9)
                    run_teacher_label.font.bold = True
                    
                    run_teacher = p1.add_run(f"{lesson['teacher']}\n")
                    run_teacher.font.size = Pt(9)
                    
                    # Время
                    run_time_label = p1.add_run(f"{headers['time']}: ")
                    run_time_label.font.size = Pt(9)
                    run_time_label.font.bold = True
                    
                    run_time = p1.add_run(f"{lesson['time']}\n")
                    run_time.font.size = Pt(9)
                    
                    # Аудитория
                    run_room_label = p1.add_run(f"{headers['room']}: ")
                    run_room_label.font.size = Pt(9)
                    run_room_label.font.bold = True
                    
                    run_room = p1.add_run(f"{lesson['classroom']}")
                    run_room.font.size = Pt(9)
                    run_room.font.color.rgb = RGBColor(192, 0, 0)
                    
                    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p1.space_after = Pt(2)
            
            # Разрыв страницы после каждой группы (кроме последней)
            if group_name != list(data_by_group.keys())[-1]:
                doc.add_page_break()
    
    # Сохранение документа
    import tempfile, os
    filename = f"schedule_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    path = os.path.join(tempfile.gettempdir(), filename)
    doc.save(path)
    
    return send_file(path, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')




@app.errorhandler(404)
def not_found(error):
    return jsonify({'error': 'Not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Internal server error'}), 500


HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>управления расписанием</title>
 <style>
        :root {
            --primary: #2563eb;
            --primary-dark: #1e40af;
            --primary-light: #3b82f6;
            --secondary: #64748b;
            --success: #10b981;
            --danger: #ef4444;
            --warning: #f59e0b;
            --bg: #f8fafc;
            --surface: #ffffff;
            --border: #e2e8f0;
            --text: #1e293b;
            --text-light: #64748b;
            --shadow: 0 1px 3px 0 rgb(0 0 0 / 0.1);
            --shadow-md: 0 4px 6px -1px rgb(0 0 0 / 0.1);
            --shadow-lg: 0 10px 15px -3px rgb(0 0 0 / 0.1);
        }

        * { margin: 0; padding: 0; box-sizing: border-box; }
        
        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            background: var(--bg);
            color: var(--text);
            min-height: 100vh;
            line-height: 1.6;
        }
        
        .container {
            max-width: 1400px;
            margin: 0 auto;
            background: var(--surface);
            border-radius: 0.75rem;
            padding: 2rem;
            box-shadow: var(--shadow-lg);
            border: 1px solid var(--border);
        }
        
        header {
            text-align: center;
            margin-bottom: 2rem;
            padding-bottom: 1.5rem;
            border-bottom: 1px solid var(--border);
        }
        
        header h1 {
            font-size: 2rem;
            color: var(--primary);
            margin-bottom: 0.5rem;
            font-weight: 600;
        }
        
        header p {
            color: var(--text-light);
            font-size: 0.95rem;
        }
        
        .tabs {
            display: flex;
            gap: 0.25rem;
            flex-wrap: wrap;
            margin-bottom: 2rem;
            background: var(--bg);
            padding: 0.375rem;
            border-radius: 0.5rem;
            border: 1px solid var(--border);
        }
        
        .tab {
            padding: 0.625rem 1rem;
            background: transparent;
            border: none;
            border-radius: 0.375rem;
            cursor: pointer;
            transition: all 0.2s;
            font-size: 0.875rem;
            font-weight: 500;
            color: var(--text-light);
        }
        
        .tab:hover { 
            background: var(--surface); 
            color: var(--text); 
        }
        
        .tab.active { 
            background: var(--primary); 
            color: white;
            box-shadow: var(--shadow);
        }
        
        .tab-content { display: none; animation: fadeIn 0.3s; }
        .tab-content.active { display: block; }
        
        @keyframes fadeIn {
            from { opacity: 0; transform: translateY(10px); }
            to { opacity: 1; transform: translateY(0); }
        }
        
        .btn {
            padding: 0.625rem 1.25rem;
            border: none;
            border-radius: 0.375rem;
            cursor: pointer;
            font-size: 0.875rem;
            font-weight: 500;
            transition: all 0.2s;
            margin: 5px;
            display: inline-flex;
            align-items: center;
            gap: 0.5rem;
        }
        
        .btn-primary { 
            background: var(--primary); 
            color: white; 
        }
        .btn-primary:hover { 
            background: var(--primary-dark); 
            transform: translateY(-1px);
            box-shadow: var(--shadow-md);
        }
        
        .btn-success { 
            background: var(--success); 
            color: white; 
        }
        .btn-success:hover { 
            background: #059669; 
        }
        
        .btn-danger { 
            background: var(--danger); 
            color: white; 
        }
        .btn-danger:hover { 
            background: #dc2626; 
        }
        
        .btn-small { 
            padding: 0.375rem 0.75rem; 
            font-size: 0.813rem; 
        }
        
        table {
            width: 100%;
            border-collapse: collapse;
            margin-top: 1.5rem;
            box-shadow: var(--shadow);
            border-radius: 0.5rem;
            overflow: hidden;
            border: 1px solid var(--border);
        }
        
        table th, table td { 
            padding: 0.875rem 1rem; 
            text-align: left; 
            border-bottom: 1px solid var(--border); 
        }
        
        table th { 
            background: var(--bg); 
            color: var(--text); 
            font-weight: 600;
            font-size: 0.813rem;
            text-transform: uppercase;
            letter-spacing: 0.025em;
        }
        
        table tr:hover { 
            background: var(--bg); 
        }
        
        table tr:last-child td {
            border-bottom: none;
        }
        
        .form-group { 
            margin-bottom: 1.25rem; 
        }
        
        .form-group label { 
            display: block; 
            margin-bottom: 0.5rem; 
            font-weight: 500; 
            color: var(--text);
            font-size: 0.875rem;
        }
        
        .form-group input, 
        .form-group select, 
        .form-group textarea {
            width: 100%;
            padding: 0.625rem 0.875rem;
            border: 1px solid var(--border);
            border-radius: 0.375rem;
            font-size: 0.875rem;
            transition: all 0.2s;
            font-family: inherit;
        }
        
        .form-group input:focus, 
        .form-group select:focus, 
        .form-group textarea:focus {
            outline: none;
            border-color: var(--primary);
            box-shadow: 0 0 0 3px rgba(37, 99, 235, 0.1);
        }
        
        .modal {
            display: none;
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: rgba(0, 0, 0, 0.5);
            z-index: 1000;
            align-items: center;
            justify-content: center;
            padding: 1rem;
        }
        .modal.show { display: flex; }
        
        .modal-content {
            background: var(--surface);
            padding: 0;
            border-radius: 0.75rem;
            max-width: 600px;
            width: 90%;
            max-height: 90vh;
            overflow-y: auto;
            box-shadow: var(--shadow-lg);
        }
        
        .modal-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            padding: 1.5rem;
            border-bottom: 1px solid var(--border);
        }
        
        .modal-header h2 { 
            color: var(--text);
            font-size: 1.25rem;
            font-weight: 600;
        }
        
        .close-modal {
            background: none;
            border: none;
            font-size: 1.75rem;
            cursor: pointer;
            color: var(--text-light);
            line-height: 1;
            padding: 0;
            width: 2rem;
            height: 2rem;
            display: flex;
            align-items: center;
            justify-content: center;
            border-radius: 0.25rem;
            transition: all 0.2s;
        }
        .close-modal:hover { 
            background: var(--bg);
            color: var(--text); 
        }
        
        .modal-content form {
            padding: 1.5rem;
        }
        
        .toast {
            position: fixed;
            bottom: 2rem;
            right: 2rem;
            background: var(--text);
            color: white;
            padding: 1rem 1.5rem;
            border-radius: 0.5rem;
            box-shadow: var(--shadow-lg);
            z-index: 2000;
            animation: slideIn 0.3s;
            max-width: 400px;
        }
        
        @keyframes slideIn {
            from { opacity: 0; transform: translateX(100px); }
            to { opacity: 1; transform: translateX(0); }
        }
        
        .toast.success { background: var(--success); }
        .toast.error { background: var(--danger); }
        .toast.warning { background: var(--warning); }
        
        .loader {
            border: 4px solid var(--bg);
            border-top: 4px solid var(--primary);
            border-radius: 50%;
            width: 50px;
            height: 50px;
            animation: spin 1s linear infinite;
            margin: 30px auto;
        }
        
        @keyframes spin {
            0% { transform: rotate(0deg); }
            100% { transform: rotate(360deg); }
        }
        
        .welcome-card {
            background: linear-gradient(135deg, var(--primary) 0%, var(--primary-dark) 100%);
            color: white;
            padding: 3rem 2rem;
            border-radius: 0.75rem;
            text-align: center;
            margin-bottom: 2rem;
        }
        
        .welcome-card h2 { 
            font-size: 2rem; 
            margin-bottom: 0.75rem; 
        }
        
        .welcome-card p { 
            font-size: 1.1rem; 
            opacity: 0.95; 
        }
        
        .stats-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 1.5rem;
            margin-top: 2rem;
        }
        
        .stat-card {
            background: var(--surface);
            padding: 1.5rem;
            border-radius: 0.5rem;
            box-shadow: var(--shadow);
            text-align: center;
            border: 1px solid var(--border);
        }
        
        .stat-card h3 { 
            font-size: 2.5rem; 
            color: var(--primary); 
            margin-bottom: 0.5rem;
            font-weight: 700;
        }
        
        .stat-card p { 
            color: var(--text-light); 
            font-size: 0.95rem; 
        }
        
        .conflict-alert {
            background: #fef2f2;
            border-left: 4px solid var(--danger);
            padding: 1rem;
            margin: 1rem 0;
            border-radius: 0.375rem;
        }
        
        .conflict-alert h4 { 
            color: var(--danger); 
            margin-bottom: 0.5rem;
            font-size: 0.95rem;
        }
        
        .conflict-alert ul {
            margin-left: 1.25rem;
            font-size: 0.875rem;
        }
        
        .checkbox-group {
            display: flex;
            align-items: center;
            gap: 0.75rem;
            margin-bottom: 1rem;
        }
        
        .checkbox-group input[type="checkbox"] {
            width: 1.125rem;
            height: 1.125rem;
            cursor: pointer;
        }
        
        .multi-select {
            min-height: 120px;
        }
        
        .teacher-stats {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 1.5rem;
            margin-top: 1.5rem;
        }
        
        .teacher-card {
            background: var(--surface);
            padding: 1.5rem;
            border-radius: 0.5rem;
            box-shadow: var(--shadow);
            border: 1px solid var(--border);
        }
        
        .teacher-card h4 {
            color: var(--primary);
            margin-bottom: 1rem;
            font-size: 1.2rem;
            font-weight: 600;
        }
        
        .time-slot {
            display: flex;
            justify-content: space-between;
            padding: 0.5rem 0;
            border-bottom: 1px solid var(--border);
        }
        
        .time-slot:last-child {
            border-bottom: none;
        }
        
        .time-slot.busy {
            background-color: #fee2e2;
        }
        
        .time-slot.free {
            background-color: #dcfce7;
        }
        
        .pair-buttons {
            display: grid;
            grid-template-columns: repeat(5, 1fr);
            gap: 0.5rem;
            margin: 1rem 0;
        }
        
        .pair-btn {
            padding: 0.75rem;
            background: var(--bg);
            border: 1px solid var(--border);
            border-radius: 0.375rem;
            cursor: pointer;
            font-weight: 500;
            transition: all 0.2s;
            font-size: 0.813rem;
        }
        
        .pair-btn:hover {
            background: var(--primary-light);
            border-color: var(--primary);
            color: white;
        }
        
        .pair-btn.active {
            background: var(--primary);
            border-color: var(--primary);
            color: white;
        }
        
        .teacher-period-stats {
            background: var(--surface);
            padding: 1.5rem;
            border-radius: 0.5rem;
            box-shadow: var(--shadow);
            margin-top: 2rem;
            border: 1px solid var(--border);
        }
        
        .teacher-period-stats h3 {
            margin-bottom: 1.5rem;
            color: var(--text);
            font-size: 1.25rem;
        }
        
        .period-selector {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 1rem;
            margin-bottom: 1.5rem;
        }
        
        .period-selector .form-group {
            margin-bottom: 0;
        }
        
        .filter-container {
            background: var(--bg);
            padding: 1rem;
            border-radius: 0.5rem;
            margin-bottom: 1.5rem;
            border: 1px solid var(--border);
            animation: slideDown 0.3s ease;
        }
        
        @keyframes slideDown {
            from {
                opacity: 0;
                transform: translateY(-10px);
            }
            to {
                opacity: 1;
                transform: translateY(0);
            }
        }
        
        .form-text {
            display: block;
            margin-top: 0.375rem;
            color: var(--text-light);
            font-size: 0.75rem;
        }
        
        .hidden { 
            display: none !important; 
        }
        
        @media (max-width: 768px) {
            .container {
                padding: 1rem;
                border-radius: 0;
            }
            
            .tabs {
                gap: 0.125rem;
            }
            
            .tab {
                padding: 0.5rem 0.75rem;
                font-size: 0.75rem;
            }
            
            header h1 {
                font-size: 1.5rem;
            }
            
            .welcome-card {
                padding: 2rem 1rem;
            }
            
            .welcome-card h2 {
                font-size: 1.5rem;
            }
            
            .stats-grid {
                grid-template-columns: 1fr;
            }
            
            .pair-buttons {
                grid-template-columns: repeat(2, 1fr);
            }
            
            .period-selector {
                grid-template-columns: 1fr;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>🎓 Система управления расписанием</h1>
            <p>Универсальная многоязычная платформа для образовательных учреждений</p>
        </header>
        
        <div class="tabs">
            <button class="tab active" onclick="switchTab('dashboard')">Панель</button>
            <button class="tab" onclick="switchTab('teachers')">Преподаватели</button>
            <button class="tab" onclick="switchTab('subjects')">Предметы</button>
            <button class="tab" onclick="switchTab('classrooms')">Аудитории</button>
            <button class="tab" onclick="switchTab('groups')">Группы</button>
            <button class="tab" onclick="switchTab('class-types')">Типы занятий</button>
            <button class="tab" onclick="switchTab('schedule')">Расписание</button>
            <button class="tab" onclick="switchTab('templates')">Шаблоны</button>
            <button class="tab" onclick="switchTab('reports')">Отчеты</button>
            <button class="tab" onclick="switchTab('settings')">Настройки</button>
        </div>
        
        <div id="dashboard" class="tab-content active">
            <div class="welcome-card">
                <h2>Добро пожаловать! 👋</h2>
                <p>Начните работу с системой, выбрав нужный раздел</p>
            </div>
            <div class="stats-grid" id="stats-container"></div>
            <div class="teacher-period-stats">
                <h3>Загруженность преподавателей по периоду</h3>
                <div class="period-selector">
                    <div class="form-group">
                        <label>Дата начала:</label>
                        <input type="date" id="teacher-stats-date-from" onchange="loadTeacherStatsByPeriod()">
                    </div>
                    <div class="form-group">
                        <label>Дата окончания:</label>
                        <input type="date" id="teacher-stats-date-to" onchange="loadTeacherStatsByPeriod()">
                    </div>
                    <div class="form-group">
                        <label>Преподаватель:</label>
                        <select id="teacher-stats-teacher" onchange="loadTeacherStatsByPeriod()">
                            <option value="">Все преподаватели</option>
                        </select>
                    </div>
                </div>
                <div id="teacher-period-stats-container"></div>
            </div>
        </div>
        
        <div id="teachers" class="tab-content">
            <h2>Управление преподавателями</h2>
            <button class="btn btn-primary" onclick="showEntityModal('teacher')">➕ Добавить</button>
            <div id="teachers-list"></div>
        </div>
        
        <div id="subjects" class="tab-content">
            <h2>Управление предметами</h2>
            <button class="btn btn-primary" onclick="showEntityModal('subject')">➕ Добавить</button>
            <div id="subjects-list"></div>
        </div>
        
        <div id="classrooms" class="tab-content">
            <h2>Управление аудиториями</h2>
            <button class="btn btn-primary" onclick="showEntityModal('classroom')">➕ Добавить</button>
            <div id="classrooms-list"></div>
        </div>
        
        <div id="groups" class="tab-content">
            <h2>Управление группами</h2>
            <button class="btn btn-primary" onclick="showEntityModal('group')">➕ Добавить</button>
            <div id="groups-list"></div>
        </div>
        
        <div id="class-types" class="tab-content">
            <h2>Типы занятий</h2>
            <button class="btn btn-primary" onclick="showEntityModal('class_type')">➕ Добавить</button>
            <div id="class-types-list"></div>
        </div>
        
        <div id="schedule" class="tab-content">
            <h2>Расписание занятий</h2>
            <div class="form-group">
                <label>Дата:</label>
                <input type="date" id="schedule-date" onchange="loadSchedule()">
            </div>
            <button class="btn btn-primary" onclick="showScheduleModal()">➕ Добавить занятие</button>
            <div id="schedule-list"></div>
        </div>
        
        <div id="templates" class="tab-content">
            <h2>Шаблоны расписания</h2>
            <button class="btn btn-primary" onclick="showTemplateModal()">📋 Создать шаблон</button>
            <div id="templates-list"></div>
        </div>
        
       <div id="reports" class="tab-content">
    <h2>Отчеты и экспорт</h2>
    
    <div class="form-group">
        <label>Язык:</label>
        <select id="report-language" class="form-control">
            <option value="russian">Русский</option>
            <option value="kyrgyz">Кыргызча</option>
            <option value="english">English</option>
        </select>
    </div>
    
    <div class="form-group">
        <label>Тип отчета:</label>
        <select id="report-type" class="form-control" onchange="onReportTypeChange()">
            <option value="general">Общий (все группы)</option>
            <option value="group">По группе/курсу</option>
            <option value="teacher">По преподавателю</option>
        </select>
    </div>

    <!-- Фильтр по группе (показывается при type="group") -->
    <div id="report-group-container" class="form-group filter-container" style="display: none;">
        <label>Выберите группу:</label>
        <select id="report-group-select" class="form-control">
            <option value="">Загрузка...</option>
        </select>
        <small class="form-text">Оставьте пустым, чтобы фильтровать только по курсу</small>
    </div>

    <!-- Фильтр по курсу (показывается при type="group") -->
    <div id="report-year-container" class="form-group filter-container" style="display: none;">
        <label>Или выберите курс:</label>
        <select id="report-year-select" class="form-control">
            <option value="">Загрузка...</option>
        </select>
        <small class="form-text">Экспорт всех групп выбранного курса</small>
    </div>

    <!-- Фильтр по преподавателю (показывается при type="teacher") -->
    <div id="report-teacher-container" class="form-group filter-container" style="display: none;">
        <label>Выберите преподавателя:</label>
        <select id="report-teacher-select" class="form-control">
            <option value="">Загрузка...</option>
        </select>
    </div>
    
    <div class="form-group">
        <label>Дата от:</label>
        <input type="date" id="report-date-from" class="form-control" required>
    </div>
    
    <div class="form-group">
        <label>Дата до:</label>
        <input type="date" id="report-date-to" class="form-control" required>
    </div>
    
    <button class="btn btn-success" onclick="exportReport()">
        📄 Экспорт DOCX
    </button>

  
</div>

<style>
.filter-container {
    background-color: #f8f9fa;
    padding: 15px;
    border-radius: 5px;

    margin-top: 10px;
    animation: slideDown 0.3s ease;
}

@keyframes slideDown {
    from {
        opacity: 0;
        transform: translateY(-10px);
    }
    to {
        opacity: 1;
        transform: translateY(0);
    }
}

.form-text {
    display: block;
    margin-top: 5px;
    color: #6c757d;
    font-size: 12px;
}

.export-info {
    border-left: 4px solid #28a745;
}

.export-info h4 {
    color: #28a745;
    font-size: 16px;
}

.export-info ul {
    list-style-type: none;
    padding-left: 0;
}

.export-info li {
    padding: 5px 0;
    padding-left: 20px;
    position: relative;
}

.export-info li:before {
    content: "✓";
    position: absolute;
    left: 0;
    color: #28a745;
    font-weight: bold;
}

/* Toast уведомления */
.toast {
    position: fixed;
    bottom: 20px;
    right: 20px;
    padding: 15px 20px;
    border-radius: 6px;
    color: white;
    font-weight: 600;
    z-index: 10000;
    animation: slideInRight 0.3s, fadeOut 0.3s 2.7s;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.3);
    min-width: 250px;
}

@keyframes slideInRight {
    from {
        transform: translateX(100%);
        opacity: 0;
    }
    to {
        transform: translateX(0);
        opacity: 1;
    }
}

@keyframes fadeOut {
    to {
        opacity: 0;
        transform: translateX(100%);
    }
}

.toast.success {
    background-color: #28a745;
}

.toast.error {
    background-color: #dc3545;
}

.toast.info {
    background-color: #17a2b8;
}

.toast.warning {
    background-color: #ffc107;
    color: #333;
}
</style>
        
        <div id="settings" class="tab-content">
            <h2>Системные настройки</h2>
            <div id="settings-form"></div>
        </div>
    </div>
    
    <div id="modal-container"></div>
    
    <script>
        let state = {
            teachers: [],
            subjects: [],
            classrooms: [],
            groups: [],
            classTypes: [],
            schedules: [],
            templates: [],
            settings: {},
            currentDate: new Date().toISOString().split('T')[0],
            editingId: null
        };
        
        document.addEventListener('DOMContentLoaded', () => {
            loadAllData();
            document.getElementById('schedule-date').value = state.currentDate;
            document.getElementById('report-date-from').value = state.currentDate;
            document.getElementById('report-date-to').value = state.currentDate;
            document.getElementById('teacher-stats-date-from').value = state.currentDate;
            document.getElementById('teacher-stats-date-to').value = state.currentDate;
        });
        
        async function loadAllData() {
            try {
                const [teachers, subjects, classrooms, groups, classTypes, settings] = await Promise.all([
                    fetch('/api/teachers').then(r => r.json()),
                    fetch('/api/subjects').then(r => r.json()),
                    fetch('/api/classrooms').then(r => r.json()),
                    fetch('/api/groups').then(r => r.json()),
                    fetch('/api/class_types').then(r => r.json()),
                    fetch('/api/settings').then(r => r.json())
                ]);
                
                state.teachers = teachers;
                state.subjects = subjects;
                state.classrooms = classrooms;
                state.groups = groups;
                state.classTypes = classTypes;
                state.settings = settings;
                
                renderAllLists();
                updateDashboardStats();
                loadTeacherStatsByPeriod();
                populateTeacherSelect();
            } catch (error) {
                showToast('Ошибка загрузки данных', 'error');
            }
        }
        
        function populateTeacherSelect() {
            const select = document.getElementById('teacher-stats-teacher');
            select.innerHTML = '<option value="">Все преподаватели</option>';
            state.teachers.forEach(teacher => {
                const option = document.createElement('option');
                option.value = teacher.id;
                option.textContent = teacher.name_russian;
                select.appendChild(option);
            });
        }
        
        function switchTab(tabName) {
            document.querySelectorAll('.tab').forEach(t => t.classList.remove('active'));
            document.querySelectorAll('.tab-content').forEach(c => c.classList.remove('active'));
            
            event.target.classList.add('active');
            document.getElementById(tabName).classList.add('active');
            
            if (tabName === 'schedule') loadSchedule();
            if (tabName === 'templates') loadTemplates();
            if (tabName === 'settings') loadSettings();
        }
        
        function updateDashboardStats() {
            const html = `
                <div class="stat-card">
                    <h3>${state.teachers.length}</h3>
                    <p>Преподавателей</p>
                </div>
                <div class="stat-card">
                    <h3>${state.subjects.length}</h3>
                    <p>Предметов</p>
                </div>
                <div class="stat-card">
                    <h3>${state.classrooms.length}</h3>
                    <p>Аудиторий</p>
                </div>
                <div class="stat-card">
                    <h3>${state.groups.length}</h3>
                    <p>Групп</p>
                </div>
                <div class="stat-card">
                    <h3>${state.schedules.length}</h3>
                    <p>Занятий сегодня</p>
                </div>
            `;
            document.getElementById('stats-container').innerHTML = html;
        }
        
        async function loadTeacherStatsByPeriod() {
            const dateFrom = document.getElementById('teacher-stats-date-from').value;
            const dateTo = document.getElementById('teacher-stats-date-to').value;
            const teacherId = document.getElementById('teacher-stats-teacher').value;
            
            if (!dateFrom || !dateTo) return;
            
            try {
                let url = `/api/schedules?date_from=${dateFrom}&date_to=${dateTo}`;
                if (teacherId) {
                    url += `&teacher_id=${teacherId}`;
                }
                
                const response = await fetch(url);
                const schedules = await response.json();
                
                // Группируем занятия по преподавателям и дням
                const teacherSchedules = {};
                schedules.forEach(schedule => {
                    if (!teacherSchedules[schedule.teacher_id]) {
                        teacherSchedules[schedule.teacher_id] = {
                            name: schedule.teacher_name_ru,
                            days: {}
                        };
                    }
                    
                    const date = schedule.date;
                    if (!teacherSchedules[schedule.teacher_id].days[date]) {
                        teacherSchedules[schedule.teacher_id].days[date] = [];
                    }
                    teacherSchedules[schedule.teacher_id].days[date].push(schedule);
                });
                
                // Сортируем по дате
                Object.values(teacherSchedules).forEach(teacher => {
                    Object.keys(teacher.days).forEach(date => {
                        teacher.days[date].sort((a, b) => a.start_time.localeCompare(b.start_time));
                    });
                });
                
                let html = '';
                
                for (const [teacherId, teacherData] of Object.entries(teacherSchedules)) {
                    html += `
                        <div class="teacher-card">
                            <h4>${teacherData.name}</h4>
                            <div>
                    `;
                    
                    // Сортируем даты
                    const sortedDates = Object.keys(teacherData.days).sort();
                    
                    sortedDates.forEach(date => {
                        const daySchedules = teacherData.days[date];
                        const dateObj = new Date(date);
                        const dayNames = ['Воскресенье', 'Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота'];
                        const dayName = dayNames[dateObj.getDay()];
                        
                        html += `<div style="margin: 10px 0; padding: 10px; background: #f8f9fa; border-radius: 6px;">
                                    <strong>${date} (${dayName})</strong>
                                    <div style="margin-top: 5px;">`;
                        
                        daySchedules.forEach(schedule => {
                            html += `<div style="padding: 5px 0; border-bottom: 1px solid #eee;">
                                        ${schedule.start_time} - ${schedule.end_time}: 
                                        ${schedule.subject_name_ru} (${schedule.group_names})
                                     </div>`;
                        });
                        
                        html += `</div></div>`;
                    });
                    
                    html += '</div></div>';
                }
                
                if (Object.keys(teacherSchedules).length === 0) {
                    html = '<p>Нет данных для отображения</p>';
                }
                
                document.getElementById('teacher-period-stats-container').innerHTML = html;
            } catch (error) {
                console.error('Ошибка загрузки статистики преподавателей:', error);
            }
        }
        
        function renderAllLists() {
            renderList('teachers', state.teachers, ['id', 'name_russian', 'name_kyrgyz', 'name_english'], 'teacher');
            renderList('subjects', state.subjects, ['id', 'name_russian', 'name_kyrgyz', 'name_english'], 'subject');
            renderList('classrooms', state.classrooms, ['id', 'name', 'capacity', 'building'], 'classroom');
            renderList('groups', state.groups, ['id', 'name', 'year', 'faculty'], 'group');
            renderList('class-types', state.classTypes, ['id', 'name_russian', 'name_kyrgyz', 'name_english'], 'class_type');
        }
        
        function renderList(entity, data, columns, entityType) {
            const container = document.getElementById(`${entity}-list`);
            if (!container) return;
            
            let html = '<table><thead><tr>';
            columns.forEach(col => html += `<th>${col}</th>`);
            html += '<th>Действия</th></tr></thead><tbody>';
            
            data.forEach(item => {
                html += '<tr>';
                columns.forEach(col => html += `<td>${item[col] || '-'}</td>`);
                html += `<td>
                    <button class="btn btn-primary btn-small" onclick="editEntity('${entityType}', ${item.id})">✏️</button>
                    <button class="btn btn-danger btn-small" onclick="deleteEntity('${entityType}', ${item.id})">🗑️</button>
                </td></tr>`;
            });
            
            html += '</tbody></table>';
            container.innerHTML = html;
        }
        
            // ENTITY CRUD MODALS
        function showEntityModal(entityType, id = null) {
            state.editingId = id;
            let title = '';
            let fields = [];
            
            if (entityType === 'teacher') {
                title = id ? 'Редактировать преподавателя' : 'Добавить преподавателя';
                fields = [
                    {name: 'name_russian', label: 'ФИО (Русский)', type: 'text', required: true},
                    {name: 'name_kyrgyz', label: 'ФИО (Кыргызча)', type: 'text'},
                    {name: 'name_english', label: 'ФИО (English)', type: 'text'}
                ];
            } else if (entityType === 'subject') {
                title = id ? 'Редактировать предмет' : 'Добавить предмет';
                fields = [
                    {name: 'name_russian', label: 'Название (Русский)', type: 'text', required: true},
                    {name: 'name_kyrgyz', label: 'Название (Кыргызча)', type: 'text'},
                    {name: 'name_english', label: 'Название (English)', type: 'text'}
                ];
            } else if (entityType === 'classroom') {
                title = id ? 'Редактировать аудиторию' : 'Добавить аудиторию';
                fields = [
                    {name: 'name', label: 'Номер/Название', type: 'text', required: true},
                    {name: 'capacity', label: 'Вместимость', type: 'number'},
                    {name: 'building', label: 'Корпус', type: 'text'}
                ];
            } else if (entityType === 'group') {
                title = id ? 'Редактировать группу' : 'Добавить группу';
                fields = [
                    {name: 'name', label: 'Название группы', type: 'text', required: true},
                    {name: 'year', label: 'Курс', type: 'number'},
                    {name: 'faculty', label: 'Факультет', type: 'text'}
                ];
            } else if (entityType === 'class_type') {
                title = id ? 'Редактировать тип занятия' : 'Добавить тип занятия';
                fields = [
                    {name: 'name_russian', label: 'Название (Русский)', type: 'text', required: true},
                    {name: 'name_kyrgyz', label: 'Название (Кыргызча)', type: 'text'},
                    {name: 'name_english', label: 'Название (English)', type: 'text'}
                ];
            }
            
            showModal(title, fields, entityType, id);
        }
        
        async function showModal(title, fields, entityType, id) {
            let formData = {};
            
            // Load existing data if editing
            if (id) {
                try {
                    const response = await fetch(`/api/${entityType === 'class_type' ? 'class_types' : entityType + 's'}/${id}`);
                    formData = await response.json();
                } catch (error) {
                    showToast('Ошибка загрузки данных', 'error');
                    return;
                }
            }
            
            let formHtml = '';
            fields.forEach(field => {
                formHtml += `
                    <div class="form-group">
                        <label>${field.label}${field.required ? ' *' : ''}:</label>
                        <input 
                            type="${field.type}" 
                            id="field-${field.name}" 
                            value="${formData[field.name] || ''}"
                            ${field.required ? 'required' : ''}
                        >
                    </div>
                `;
            });
            
            const modalHtml = `
                <div class="modal show" id="entity-modal">
                    <div class="modal-content">
                        <div class="modal-header">
                            <h2>${title}</h2>
                            <button class="close-modal" onclick="closeModal()">&times;</button>
                        </div>
                        <form id="entity-form" onsubmit="saveEntity(event, '${entityType}', ${id})">
                            ${formHtml}
                            <div style="margin-top: 20px;">
                                <button type="submit" class="btn btn-success">💾 Сохранить</button>
                                <button type="button" class="btn btn-danger" onclick="closeModal()">❌ Отмена</button>
                            </div>
                        </form>
                    </div>
                </div>
            `;
            
            document.getElementById('modal-container').innerHTML = modalHtml;
        }
        
        async function saveEntity(event, entityType, id) {
            event.preventDefault();
            
            const form = event.target;
            const formData = new FormData(form);
            const data = {};
            
            // Collect all input values
            form.querySelectorAll('input, select, textarea').forEach(input => {
                if (input.id.startsWith('field-')) {
                    const fieldName = input.id.replace('field-', '');
                    data[fieldName] = input.type === 'number' ? parseInt(input.value) || null : input.value;
                }
            });
            
            try {
                const endpoint = entityType === 'class_type' ? 'class_types' : entityType + 's';
                const url = id ? `/api/${endpoint}/${id}` : `/api/${endpoint}`;
                const method = id ? 'PUT' : 'POST';
                
                const response = await fetch(url, {
                    method: method,
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(data)
                });
                
                if (response.ok) {
                    showToast('✅ Успешно сохранено', 'success');
                    closeModal();
                    await loadAllData();
                } else {
                    const error = await response.json();
                    showToast(`❌ ${error.error || 'Ошибка сохранения'}`, 'error');
                }
            } catch (error) {
                showToast('❌ Ошибка сохранения', 'error');
            }
        }
        
        async function editEntity(entityType, id) {
            showEntityModal(entityType, id);
        }
        
        async function deleteEntity(entityType, id) {
            if (!confirm('Вы уверены, что хотите удалить?')) return;
            
            try {
                const endpoint = entityType === 'class_type' ? 'class_types' : entityType + 's';
                const response = await fetch(`/api/${endpoint}/${id}`, {
                    method: 'DELETE'
                });
                
                if (response.ok) {
                    showToast('✅ Успешно удалено', 'success');
                    await loadAllData();
                } else {
                    const error = await response.json();
                    showToast(`❌ ${error.error || 'Ошибка удаления'}`, 'error');
                }
            } catch (error) {
                showToast('❌ Ошибка удаления', 'error');
            }
        }
        
        // SCHEDULE MANAGEMENT
async function showScheduleModal(id = null) {
    state.editingId = id;
    let scheduleData = {
        date: document.getElementById('schedule-date').value,
        start_time: '09:00',
        end_time: '10:30',
        teacher_id: '',
        subject_id: '',
        classroom_id: '',
        class_type_id: '',
        is_alternating: 0,
        week_type: '',
        break_after_minutes: state.settings.default_break_minutes || 10,
        notes: '',
        group_ids: []
    };
    
    if (id) {
        try {
            const response = await fetch(`/api/schedules/${id}`);
            const data = await response.json();
            scheduleData = {...scheduleData, ...data};
        } catch (error) {
            showToast('Ошибка загрузки данных', 'error');
            return;
        }
    }
    
    // Запрос типа недели
    const weekTypeResp = await fetch(`/api/week_type?date=${scheduleData.date}`);
    const weekTypeData = await weekTypeResp.json();
    const calculatedWeekType = weekTypeData.week_type;
    
    // Генерируем кнопки для пар
    let pairButtonsHtml = '';
    for (let i = 1; i <= 10; i++) {
        const startTimeKey = `start_time_${i}`;
        const startTime = state.settings[startTimeKey];
        if (startTime) {
            pairButtonsHtml += `<button type="button" class="pair-btn" onclick="setPairTime(${i})">${i} пара (${startTime})</button>`;
        }
    }
    
    const modalHtml = `
        <div class="modal show" id="schedule-modal">
            <div class="modal-content">
                <div class="modal-header">
                    <h2>${id ? 'Редактировать' : 'Добавить'} занятие</h2>
                    <button class="close-modal" onclick="closeModal()">&times;</button>
                </div>
                <form id="schedule-form" onsubmit="saveSchedule(event, ${id})">
                    <div class="form-group">
                        <label>Дата *:</label>
                        <input type="date" id="field-date" value="${scheduleData.date}" required>
                    </div>
                    
                    <div class="form-group">
                        <label>Быстрый выбор пары:</label>
                        <div class="pair-buttons">
                            ${pairButtonsHtml}
                        </div>
                    </div>
                    
                    <div class="form-group">
                        <label>Время начала *:</label>
                        <input type="time" id="field-start_time" value="${scheduleData.start_time}" required>
                    </div>
                    
                    <div class="form-group">
                        <label>Время окончания *:</label>
                        <input type="time" id="field-end_time" value="${scheduleData.end_time}" required>
                    </div>
                    
                    <div class="form-group">
                        <label>Преподаватель *:</label>
                        <select id="field-teacher_id" required>
                            <option value="">Выберите...</option>
                            ${state.teachers.map(t => `
                                <option value="${t.id}" ${t.id == scheduleData.teacher_id ? 'selected' : ''}>
                                    ${t.name_russian}
                                </option>
                            `).join('')}
                        </select>
                    </div>
                    
                    <div class="form-group">
                        <label>Предмет *:</label>
                        <select id="field-subject_id" required>
                            <option value="">Выберите...</option>
                            ${state.subjects.map(s => `
                                <option value="${s.id}" ${s.id == scheduleData.subject_id ? 'selected' : ''}>
                                    ${s.name_russian}
                                </option>
                            `).join('')}
                        </select>
                    </div>
                    
                    <div class="form-group">
                        <label>Аудитория *:</label>
                        <select id="field-classroom_id" required>
                            <option value="">Выберите...</option>
                            ${state.classrooms.map(c => `
                                <option value="${c.id}" ${c.id == scheduleData.classroom_id ? 'selected' : ''}>
                                    ${c.name}
                                </option>
                            `).join('')}
                        </select>
                    </div>
                    
                    <div class="form-group">
                        <label>Тип занятия *:</label>
                        <select id="field-class_type_id" required>
                            <option value="">Выберите...</option>
                            ${state.classTypes.map(ct => `
                                <option value="${ct.id}" ${ct.id == scheduleData.class_type_id ? 'selected' : ''}>
                                    ${ct.name_russian}
                                </option>
                            `).join('')}
                        </select>
                    </div>
                    
                    <div class="form-group">
                        <label>Группы *:</label>
                        <select id="field-group_ids" multiple class="multi-select" required>
                            ${state.groups.map(g => `
                                <option value="${g.id}" ${scheduleData.group_ids && scheduleData.group_ids.includes(g.id) ? 'selected' : ''}>
                                    ${g.name}
                                </option>
                            `).join('')}
                        </select>
                        <small>Удерживайте Ctrl для множественного выбора</small>
                    </div>
                    
                    <div class="form-group checkbox-group">
                        <input 
                            type="checkbox" 
                            id="field-is_alternating" 
                            ${scheduleData.is_alternating ? 'checked' : ''}
                            onchange="toggleWeekType()"
                        >
                        <label for="field-is_alternating">Чередующийся предмет (числитель/знаменатель)</label>
                    </div>
                    
                    <div class="form-group" id="week-type-group" style="display: ${scheduleData.is_alternating ? 'block' : 'none'};">
                        <label>Тип недели:</label>
                        <select id="field-week_type">
                            <option value="numerator" ${calculatedWeekType === 'numerator' ? 'selected' : ''}>Числитель</option>
                            <option value="denominator" ${calculatedWeekType === 'denominator' ? 'selected' : ''}>Знаменатель</option>
                        </select>
                        <small>Автоматически: ${calculatedWeekType === 'numerator' ? 'Числитель' : 'Знаменатель'}</small>
                    </div>
                    
                    <div class="form-group">
                        <label>Перерыв после занятия (минут):</label>
                        <input type="number" id="field-break_after_minutes" value="${scheduleData.break_after_minutes || 10}">
                    </div>
                    
                    <div class="form-group">
                        <label>Примечания:</label>
                        <textarea id="field-notes" rows="3">${scheduleData.notes || ''}</textarea>
                    </div>
                    
                    <div id="conflict-container"></div>
                    
                    <div style="margin-top: 20px;">
                        <button type="submit" class="btn btn-success">💾 Сохранить</button>
                        <button type="button" class="btn btn-primary" onclick="validateScheduleForm()">🔍 Проверить конфликты</button>
                        <button type="button" class="btn btn-danger" onclick="closeModal()">❌ Отмена</button>
                    </div>
                </form>
            </div>
        </div>
    `;
    
    document.getElementById('modal-container').innerHTML = modalHtml;
    
    const startTimeInput = document.getElementById('field-start_time');
    const endTimeInput = document.getElementById('field-end_time');
    
    // Получаем текущие системные настройки
    const duration = state.settings.class_duration_minutes || 80;
    
    // Автоматическое обновление времени конца по изменению начала
    startTimeInput.addEventListener('change', function () {
        if (this.value) {
            const [hours, minutes] = this.value.split(':').map(Number);
            const startDate = new Date();
            startDate.setHours(hours, minutes, 0);
            const endDate = new Date(startDate.getTime() + duration * 60000);
            const endHours = String(endDate.getHours()).padStart(2, '0');
            const endMinutes = String(endDate.getMinutes()).padStart(2, '0');
            endTimeInput.value = `${endHours}:${endMinutes}`;
        }
    });
}

// Функция для установки времени пары по кнопке
function setPairTime(pairNumber) {
    const startTimeKey = `start_time_${pairNumber}`;
    const startTime = state.settings[startTimeKey];
    
    if (startTime) {
        document.getElementById('field-start_time').value = startTime;
        
        // Автоматически устанавливаем время окончания
        const duration = state.settings.class_duration_minutes || 80;
        const [hours, minutes] = startTime.split(':').map(Number);
        const startDate = new Date();
        startDate.setHours(hours, minutes, 0);
        const endDate = new Date(startDate.getTime() + duration * 60000);
        const endHours = String(endDate.getHours()).padStart(2, '0');
        const endMinutes = String(endDate.getMinutes()).padStart(2, '0');
        document.getElementById('field-end_time').value = `${endHours}:${endMinutes}`;
        
        // Подсвечиваем активную кнопку
        document.querySelectorAll('.pair-btn').forEach(btn => btn.classList.remove('active'));
        event.target.classList.add('active');
    }
}

        
      function collectScheduleFormData() {
    const groupSelect = document.getElementById('field-group_ids');
    const selectedGroups = Array.from(groupSelect.selectedOptions).map(opt => parseInt(opt.value));
    
    const startTime = document.getElementById('field-start_time').value;
    const endTime = document.getElementById('field-end_time').value;
    
    return {
        id: state.editingId,
        date: document.getElementById('field-date').value,
        start_time: startTime,
        end_time: endTime,
        teacher_id: parseInt(document.getElementById('field-teacher_id').value),
        subject_id: parseInt(document.getElementById('field-subject_id').value),
        classroom_id: parseInt(document.getElementById('field-classroom_id').value),
        class_type_id: parseInt(document.getElementById('field-class_type_id').value),
        is_alternating: document.getElementById('field-is_alternating').checked ? 1 : 0,
        week_type: document.getElementById('field-is_alternating').checked ? 
                   document.getElementById('field-week_type').value : null,
        break_after_minutes: parseInt(document.getElementById('field-break_after_minutes').value) || null,
        notes: document.getElementById('field-notes').value,
        group_ids: selectedGroups
    };
}

        
        function displayConflicts(conflicts) {
            const container = document.getElementById('conflict-container');
            
            if (!conflicts || conflicts.length === 0) {
                container.innerHTML = '';
                return;
            }
            
            let html = '<div class="conflict-alert"><h4>⚠️ Обнаружены конфликты:</h4><ul>';
            conflicts.forEach(c => {
                html += `<li>${c.message}</li>`;
            });
            html += '</ul></div>';
            
            container.innerHTML = html;
        }
        
        async function saveSchedule(event, id) {
            event.preventDefault();
            
            const data = collectScheduleFormData();
            
            try {
                const url = id ? `/api/schedules/${id}` : '/api/schedules';
                const method = id ? 'PUT' : 'POST';
                
                const response = await fetch(url, {
                    method: method,
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(data)
                });
                
                if (response.ok) {
                    showToast('✅ Расписание сохранено', 'success');
                    closeModal();
                    await loadSchedule();
                } else if (response.status === 409) {
                    const error = await response.json();
                    displayConflicts(error.conflicts);
                    showToast('❌ Обнаружены конфликты', 'error');
                } else {
                    showToast('❌ Ошибка сохранения', 'error');
                }
            } catch (error) {
                showToast('❌ Ошибка сохранения', 'error');
            }
        }
        
        async function loadSchedule() {
            const date = document.getElementById('schedule-date').value;
            
            try {
                const response = await fetch(`/api/schedules?date=${date}`);
                state.schedules = await response.json();
                renderSchedule();
            } catch (error) {
                showToast('Ошибка загрузки расписания', 'error');
            }
        }
        
        function renderSchedule() {
            const container = document.getElementById('schedule-list');
            
            if (state.schedules.length === 0) {
                container.innerHTML = '<p style="margin-top: 20px;">Нет расписания на выбранную дату</p>';
                return;
            }
            
            let html = '<table><thead><tr>';
            html += '<th>Время</th><th>Группа</th><th>Предмет</th><th>Преподаватель</th>';
            html += '<th>Аудитория</th><th>Тип</th><th>Неделя</th><th>Действия</th>';
            html += '</tr></thead><tbody>';
            
            state.schedules.forEach(s => {
                const weekTypeLabel = s.is_alternating ? 
                    (s.week_type === 'numerator' ? '🔢 Числитель' : '🔤 Знаменатель') : 
                    '—';
                
                html += `<tr>
                    <td>${s.start_time} - ${s.end_time}</td>
                    <td>${s.group_names || '-'}</td>
                    <td>${s.subject_name_ru}</td>
                    <td>${s.teacher_name_ru}</td>
                    <td>${s.classroom_name}</td>
                    <td>${s.class_type_ru}</td>
                    <td>${weekTypeLabel}</td>
                    <td>
                        <button class="btn btn-primary btn-small" onclick="showScheduleModal(${s.id})">✏️</button>
                        <button class="btn btn-danger btn-small" onclick="deleteSchedule(${s.id})">🗑️</button>
                    </td>
                </tr>`;
            });
            
            html += '</tbody></table>';
            container.innerHTML = html;
        }
        
        async function deleteSchedule(id) {
            if (!confirm('Удалить занятие?')) return;
            
            try {
                const response = await fetch(`/api/schedules/${id}`, {
                    method: 'DELETE'
                });
                
                if (response.ok) {
                    showToast('✅ Занятие удалено', 'success');
                    await loadSchedule();
                } else {
                    showToast('❌ Ошибка удаления', 'error');
                }
            } catch (error) {
                showToast('❌ Ошибка удаления', 'error');
            }
        }
        
        // TEMPLATES
        async function showTemplateModal() {
            const modalHtml = `
                <div class="modal show" id="template-modal">
                    <div class="modal-content">
                        <div class="modal-header">
                            <h2>Создать шаблон</h2>
                            <button class="close-modal" onclick="closeModal()">&times;</button>
                        </div>
                        <form id="template-form" onsubmit="createTemplate(event)">
                            <div class="form-group">
                                <label>Название шаблона *:</label>
                                <input type="text" id="field-template-name" required>
                            </div>
                            
                            <div class="form-group">
                                <label>Описание:</label>
                                <textarea id="field-template-description" rows="3"></textarea>
                            </div>
                            
                            <div class="form-group">
                                <label>Дата источника *:</label>
                                <input type="date" id="field-template-date" value="${state.currentDate}" required>
                                <small>Выберите дату, с которой нужно скопировать расписание</small>
                            </div>
                            
                            <div style="margin-top: 20px;">
                                <button type="submit" class="btn btn-success">💾 Создать шаблон</button>
                                <button type="button" class="btn btn-danger" onclick="closeModal()">❌ Отмена</button>
                            </div>
                        </form>
                    </div>
                </div>
            `;
            
            document.getElementById('modal-container').innerHTML = modalHtml;
        }
        
        async function createTemplate(event) {
            event.preventDefault();
            
            const data = {
                name: document.getElementById('field-template-name').value,
                description: document.getElementById('field-template-description').value,
                source_date: document.getElementById('field-template-date').value
            };
            
            try {
                const response = await fetch('/api/templates', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(data)
                });
                
                if (response.ok) {
                    showToast('✅ Шаблон создан', 'success');
                    closeModal();
                    await loadTemplates();
                } else {
                    showToast('❌ Ошибка создания шаблона', 'error');
                }
            } catch (error) {
                showToast('❌ Ошибка создания шаблона', 'error');
            }
        }
        
        async function loadTemplates() {
            try {
                const response = await fetch('/api/templates');
                state.templates = await response.json();
                renderTemplates();
            } catch (error) {
                showToast('Ошибка загрузки шаблонов', 'error');
            }
        }
        
        function renderTemplates() {
            const container = document.getElementById('templates-list');
            
            if (state.templates.length === 0) {
                container.innerHTML = '<p style="margin-top: 20px;">Нет созданных шаблонов</p>';
                return;
            }
            
            let html = '<table><thead><tr>';
            html += '<th>ID</th><th>Название</th><th>Описание</th><th>Дата источника</th><th>Создан</th><th>Действия</th>';
            html += '</tr></thead><tbody>';
            
            state.templates.forEach(t => {
                html += `<tr>
                    <td>${t.id}</td>
                    <td>${t.name}</td>
                    <td>${t.description || '-'}</td>
                    <td>${t.source_date || '-'}</td>
                    <td>${t.created_at ? new Date(t.created_at).toLocaleDateString() : '-'}</td>
                    <td>
                        <button class="btn btn-primary btn-small" onclick="showApplyTemplateModal(${t.id})">📋 Применить</button>
                        <button class="btn btn-danger btn-small" onclick="deleteTemplate(${t.id})">🗑️</button>
                    </td>
                </tr>`;
            });
            
            html += '</tbody></table>';
            container.innerHTML = html;
        }
        
        async function showApplyTemplateModal(templateId) {
            const modalHtml = `
                <div class="modal show" id="apply-template-modal">
                    <div class="modal-content">
                        <div class="modal-header">
                            <h2>Применить шаблон</h2>
                            <button class="close-modal" onclick="closeModal()">&times;</button>
                        </div>
                        <form id="apply-template-form" onsubmit="applyTemplate(event, ${templateId})">
                            <div class="form-group">
                                <label>Дата начала *:</label>
                                <input type="date" id="field-apply-date-from" required>
                            </div>
                            
                            <div class="form-group">
                                <label>Дата окончания *:</label>
                                <input type="date" id="field-apply-date-to" required>
                            </div>
                            
                            <div style="margin-top: 20px;">
                                <button type="submit" class="btn btn-success">✅ Применить</button>
                                <button type="button" class="btn btn-danger" onclick="closeModal()">❌ Отмена</button>
                            </div>
                        </form>
                    </div>
                </div>
            `;
            
            document.getElementById('modal-container').innerHTML = modalHtml;
        }
        
        async function applyTemplate(event, templateId) {
            event.preventDefault();
            
            const data = {
                date_range: {
                    start: document.getElementById('field-apply-date-from').value,
                    end: document.getElementById('field-apply-date-to').value
                }
            };
            
            try {
                const response = await fetch(`/api/templates/${templateId}/apply`, {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(data)
                });
                
                const result = await response.json();
                
                if (result.success) {
                    showToast(`✅ Создано занятий: ${result.created_count}`, 'success');
                    
                    if (result.conflicts && result.conflicts.length > 0) {
                        showToast(`⚠️ Найдено конфликтов: ${result.conflicts.length}`, 'warning');
                    }
                    
                    closeModal();
                    await loadSchedule();
                } else {
                    showToast('❌ Ошибка применения шаблона', 'error');
                }
            } catch (error) {
                showToast('❌ Ошибка применения шаблона', 'error');
            }
        }
        
        async function deleteTemplate(id) {
            if (!confirm('Удалить шаблон?')) return;
            
            try {
                const response = await fetch(`/api/templates/${id}`, {
                    method: 'DELETE'
                });
                
                if (response.ok) {
                    showToast('✅ Шаблон удален', 'success');
                    await loadTemplates();
                } else {
                    showToast('❌ Ошибка удаления', 'error');
                }
            } catch (error) {
                showToast('❌ Ошибка удаления', 'error');
            }
        }
        
        // SETTINGS
        async function loadSettings() {
            const form = document.getElementById('settings-form');
            
            form.innerHTML = `
                <div class="form-group">
                    <label>Длительность занятия (минуты):</label>
                    <input type="number" id="setting-class-duration" value="${state.settings.class_duration_minutes || 80}">
                </div>
                
                <div class="form-group">
                    <label>Перерыв по умолчанию (минуты):</label>
                    <input type="number" id="setting-default-break" value="${state.settings.default_break_minutes || 10}">
                </div>
                
                <div class="form-group">
                    <label>Время начала 1-й пары:</label>
                    <input type="time" id="setting-start-time-1" value="${state.settings.start_time_1 || '08:00'}">
                </div>
                
                <div class="form-group">
                    <label>Время начала 2-й пары:</label>
                    <input type="time" id="setting-start-time-2" value="${state.settings.start_time_2 || '09:30'}">
                </div>
                
                <div class="form-group">
                    <label>Время начала 3-й пары:</label>
                    <input type="time" id="setting-start-time-3" value="${state.settings.start_time_3 || '11:10'}">
                </div>
                
                <div class="form-group">
                    <label>Время начала 4-й пары:</label>
                    <input type="time" id="setting-start-time-4" value="${state.settings.start_time_4 || '12:50'}">
                </div>
                
                <div class="form-group">
                    <label>Время начала 5-й пары:</label>
                    <input type="time" id="setting-start-time-5" value="${state.settings.start_time_5 || '14:30'}">
                </div>
                
                <div class="form-group">
                    <label>Время начала 6-й пары:</label>
                    <input type="time" id="setting-start-time-6" value="${state.settings.start_time_6 || '16:10'}">
                </div>
                
                <div class="form-group">
                    <label>Время начала 7-й пары:</label>
                    <input type="time" id="setting-start-time-7" value="${state.settings.start_time_7 || '17:50'}">
                </div>
                
                <div class="form-group">
                    <label>Время начала 8-й пары:</label>
                    <input type="time" id="setting-start-time-8" value="${state.settings.start_time_8 || '19:30'}">
                </div>
                
                <div class="form-group">
                    <label>Время начала 9-й пары:</label>
                    <input type="time" id="setting-start-time-9" value="${state.settings.start_time_9 || '21:10'}">
                </div>
                
                <div class="form-group">
                    <label>Время начала 10-й пары:</label>
                    <input type="time" id="setting-start-time-10" value="${state.settings.start_time_10 || '22:50'}">
                </div>
                
                <div class="form-group">
                    <label>Текущий тип недели:</label>
                    <select id="setting-current-week-type">
                        <option value="numerator" ${state.settings.current_week_type === 'numerator' ? 'selected' : ''}>Числитель</option>
                        <option value="denominator" ${state.settings.current_week_type === 'denominator' ? 'selected' : ''}>Знаменатель</option>
                    </select>
                </div>
                
                <div class="form-group">
                    <label>Дата начала отсчета недель:</label>
                    <input type="date" id="setting-week-start-date" value="${state.settings.week_start_date || state.currentDate}">
                </div>
                
                <div class="form-group">
                    <label>Рабочих дней в неделю:</label>
                    <input type="number" id="setting-working-days" value="${state.settings.working_days_per_week || 6}" min="1" max="7">
                </div>
                
                <button class="btn btn-success" onclick="saveSettings()">💾 Сохранить настройки</button>
            `;
        }
        
        async function saveSettings() {
            const data = {
                class_duration_minutes: parseInt(document.getElementById('setting-class-duration').value),
                default_break_minutes: parseInt(document.getElementById('setting-default-break').value),
                start_time_1: document.getElementById('setting-start-time-1').value,
                start_time_2: document.getElementById('setting-start-time-2').value,
                start_time_3: document.getElementById('setting-start-time-3').value,
                start_time_4: document.getElementById('setting-start-time-4').value,
                start_time_5: document.getElementById('setting-start-time-5').value,
                start_time_6: document.getElementById('setting-start-time-6').value,
                start_time_7: document.getElementById('setting-start-time-7').value,
                start_time_8: document.getElementById('setting-start-time-8').value,
                start_time_9: document.getElementById('setting-start-time-9').value,
                start_time_10: document.getElementById('setting-start-time-10').value,
                current_week_type: document.getElementById('setting-current-week-type').value,
                week_start_date: document.getElementById('setting-week-start-date').value,
                working_days_per_week: parseInt(document.getElementById('setting-working-days').value)
            };
            
            try {
                const response = await fetch('/api/settings', {
                    method: 'PUT',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify(data)
                });
                
                if (response.ok) {
                    showToast('✅ Настройки сохранены', 'success');
                    await loadAllData();
                } else {
                    showToast('❌ Ошибка сохранения', 'error');
                }
            } catch (error) {
                showToast('❌ Ошибка сохранения', 'error');
            }
        }
        
        // EXPORT
// ---------- EXPORT.RAPORT.JS ----------

// Функция определения первого дня недели (понедельника)
function getMonday(date) {
  const d = new Date(date);
  const day = d.getDay();
  const diff = (day === 0 ? -6 : 1) - day; // если воскресенье (0), то -6 иначе разница до понедельника
  d.setDate(d.getDate() + diff);
  return d.toISOString().split('T')[0];
}

// Функция для загрузки групп в селект
async function loadGroupsForReport() {
  try {
    const response = await fetch('/api/groups');
    const groups = await response.json();
    
    const groupSelect = document.getElementById('report-group-select');
    if (groupSelect) {
      groupSelect.innerHTML = '<option value="">Все группы</option>';
      groups.forEach(group => {
        const option = document.createElement('option');
        option.value = group.id;
        option.textContent = `${group.name}${group.year ? ` (${group.year} курс)` : ''}`;
        groupSelect.appendChild(option);
      });
    }
  } catch (error) {
    console.error('Ошибка загрузки групп:', error);
  }
}

// Функция для загрузки курсов в селект
async function loadYearsForReport() {
  try {
    const response = await fetch('/api/groups');
    const groups = await response.json();
    
    // Получаем уникальные курсы
    const years = [...new Set(groups.map(g => g.year).filter(y => y !== null))].sort();
    
    const yearSelect = document.getElementById('report-year-select');
    if (yearSelect) {
      yearSelect.innerHTML = '<option value="">Все курсы</option>';
      years.forEach(year => {
        const option = document.createElement('option');
        option.value = year;
        option.textContent = `${year} курс`;
        yearSelect.appendChild(option);
      });
    }
  } catch (error) {
    console.error('Ошибка загрузки курсов:', error);
  }
}

// Функция для загрузки преподавателей в селект
async function loadTeachersForReport() {
  try {
    const response = await fetch('/api/teachers');
    const teachers = await response.json();
    
    const teacherSelect = document.getElementById('report-teacher-select');
    if (teacherSelect) {
      teacherSelect.innerHTML = '<option value="">Все преподаватели</option>';
      teachers.forEach(teacher => {
        const option = document.createElement('option');
        option.value = teacher.id;
        option.textContent = teacher.name_russian;
        teacherSelect.appendChild(option);
      });
    }
  } catch (error) {
    console.error('Ошибка загрузки преподавателей:', error);
  }
}

// Обработчик изменения типа отчета
function onReportTypeChange() {
  const reportType = document.getElementById('report-type').value;
  
  // Скрываем все дополнительные фильтры
  const groupContainer = document.getElementById('report-group-container');
  const yearContainer = document.getElementById('report-year-container');
  const teacherContainer = document.getElementById('report-teacher-container');
  
  if (groupContainer) groupContainer.style.display = 'none';
  if (yearContainer) yearContainer.style.display = 'none';
  if (teacherContainer) teacherContainer.style.display = 'none';
  
  // Показываем нужный фильтр
  if (reportType === 'group') {
    if (groupContainer) groupContainer.style.display = 'block';
    if (yearContainer) yearContainer.style.display = 'block';
    loadGroupsForReport();
    loadYearsForReport();
  } else if (reportType === 'teacher') {
    if (teacherContainer) teacherContainer.style.display = 'block';
    loadTeachersForReport();
  }
}

// Функция определения первого дня недели (понедельника)
function getMonday(date) {
  const d = new Date(date);
  const day = d.getDay();
  const diff = (day === 0 ? -6 : 1) - day;
  d.setDate(d.getDate() + diff);
  return d.toISOString().split('T')[0];
}

// Обновлённая функция экспорта отчёта
async function exportReport() {
  const language = document.getElementById('report-language').value;
  const reportType = document.getElementById('report-type').value;
  let dateFrom = document.getElementById('report-date-from').value;
  let dateTo = document.getElementById('report-date-to').value;

  // Проверка на заполненность дат
  if (!dateFrom || !dateTo) {
    showToast('Пожалуйста, укажите период экспорта', 'warning');
    return;
  }

  // Подготовка фильтров
  const filters = {
    date_from: dateFrom,
    date_to: dateTo,
  };

  // Добавляем специфичные фильтры в зависимости от типа отчета
  if (reportType === 'group') {
    const groupId = document.getElementById('report-group-select')?.value;
    const year = document.getElementById('report-year-select')?.value;
    
    if (groupId) {
      filters.group_id = parseInt(groupId);
    }
    if (year && !groupId) {
      // Если выбран курс, но не выбрана конкретная группа
      filters.year = parseInt(year);
    }
  } else if (reportType === 'teacher') {
    const teacherId = document.getElementById('report-teacher-select')?.value;
    if (teacherId) {
      filters.teacher_id = parseInt(teacherId);
    }
  }

  try {
    const response = await fetch('/api/export_docx', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({
        language,
        report_type: reportType,
        filters: filters,
      }),
    });

    if (response.ok) {
      const blob = await response.blob();
      const url = window.URL.createObjectURL(blob);
      const a = document.createElement('a');
      a.href = url;

      // Формирование имени файла
      let filename = `schedule_${reportType}_${dateFrom}_${dateTo}`;
      
      if (reportType === 'group') {
        const groupSelect = document.getElementById('report-group-select');
        if (groupSelect && groupSelect.value) {
          const groupName = groupSelect.options[groupSelect.selectedIndex].text;
          filename += `_${groupName.replace(/\s+/g, '_').replace(/[()]/g, '')}`;
        } else {
          const yearSelect = document.getElementById('report-year-select');
          if (yearSelect && yearSelect.value) {
            filename += `_${yearSelect.value}_course`;
          }
        }
      } else if (reportType === 'teacher') {
        const teacherSelect = document.getElementById('report-teacher-select');
        if (teacherSelect && teacherSelect.value) {
          const teacherName = teacherSelect.options[teacherSelect.selectedIndex].text;
          filename += `_${teacherName.replace(/\s+/g, '_')}`;
        }
      }
      
      a.download = `${filename}.docx`;
      a.click();
      
      // Освобождаем память
      window.URL.revokeObjectURL(url);

      showToast('Экспорт успешно завершён', 'success');
    } else {
      const error = await response.json();
      showToast(error.error || 'Ошибка при экспорте', 'error');
    }
  } catch (error) {
    console.error('Ошибка экспорта:', error);
    showToast('Ошибка подключения при экспорте', 'error');
  }
}

// Функция для показа уведомлений
function showToast(message, type = 'info') {
  // Удаляем предыдущие toast, если есть
  const existingToast = document.querySelector('.toast');
  if (existingToast) {
    existingToast.remove();
  }

  const toast = document.createElement('div');
  toast.className = `toast ${type}`;
  toast.textContent = message;
  document.body.appendChild(toast);

  // Автоматическое удаление через 3 секунды
  setTimeout(() => {
    toast.remove();
  }, 3000);
}

// Инициализация при загрузке страницы
document.addEventListener('DOMContentLoaded', function() {
  // Устанавливаем обработчик изменения типа отчета
  const reportTypeSelect = document.getElementById('report-type');
  if (reportTypeSelect) {
    reportTypeSelect.addEventListener('change', onReportTypeChange);
  }

  // Устанавливаем текущую дату по умолчанию
  const today = new Date().toISOString().split('T')[0];
  const dateFromInput = document.getElementById('report-date-from');
  const dateToInput = document.getElementById('report-date-to');
  
  if (dateFromInput && !dateFromInput.value) {
    dateFromInput.value = today;
  }
  if (dateToInput && !dateToInput.value) {
    dateToInput.value = today;
  }
});
        
        // UTILITY FUNCTIONS
        function closeModal() {
            document.getElementById('modal-container').innerHTML = '';
        }
        
        function showToast(message, type = 'info') {
            const toast = document.createElement('div');
            toast.className = `toast ${type}`;
            toast.textContent = message;
            document.body.appendChild(toast);
            
            setTimeout(() => {
                toast.style.animation = 'slideIn 0.3s reverse';
                setTimeout(() => toast.remove(), 300);
            }, 3000);
        }
    </script>
</body>
</html>
"""

# ==================== MAIN ====================

if __name__ == '__main__':
    print("=" * 60)
    print("🎓 СИСТЕМА УПРАВЛЕНИЯ РАСПИСАНИЕМ")
    print("=" * 60)

    print("🌐 Запуск сервера...")
    print("=" * 60)
    print("📍 Открой в браузере: http://localhost:5000")
    print("=" * 60)
    print("=" * 60)
    
    app.run(debug=True, host='0.0.0.0', port=5000, threaded=True)
